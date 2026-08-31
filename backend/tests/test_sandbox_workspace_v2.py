import hashlib
import sys
from pathlib import Path

import pytest
from docx import Document

from app.cowork.sandbox import (
    SandboxLimits,
    _commit_outputs,
    build_sandbox_launch,
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_sandbox_outputs_are_validated_and_never_overwrite_without_baseline(
    tmp_path: Path,
) -> None:
    outputs = tmp_path / "container-outputs"
    destination = tmp_path / "workspace"
    outputs.mkdir()
    destination.mkdir()
    (destination / "existing.txt").write_text("original", encoding="utf-8")
    (outputs / "existing.txt").write_text("replacement", encoding="utf-8")
    (outputs / "safe.txt").write_text("safe", encoding="utf-8")
    (outputs / "fallback.pptx").write_bytes(b"not-a-presentation")
    (outputs / "active.html").write_text(
        "<!doctype html><script>alert(1)</script>", encoding="utf-8"
    )

    committed, warnings = _commit_outputs(
        outputs,
        destination,
        max_files=10,
        max_bytes=1024 * 1024,
    )

    assert committed == (str(destination / "safe.txt"),)
    assert (destination / "existing.txt").read_text(encoding="utf-8") == "original"
    assert not (destination / "active.html").exists()
    assert any("不允许无 baseline 覆盖" in warning for warning in warnings)
    assert any("候选产物验证失败" in warning for warning in warnings)
    assert any("必须使用 render_artifact" in warning for warning in warnings)


def test_linux_sandbox_never_mounts_host_root(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.cowork.sandbox.sys.platform", "linux")
    monkeypatch.setattr(
        "app.cowork.sandbox.shutil.which",
        lambda name: "/usr/bin/bwrap" if name == "bwrap" else None,
    )
    directories = {
        name: tmp_path / name for name in ("inputs", "work", "outputs", "tmp", "runtime")
    }
    for directory in directories.values():
        directory.mkdir()

    launch = build_sandbox_launch(
        command="$WORKPILOT_PYTHON build.py",
        inputs=directories["inputs"],
        work=directories["work"],
        outputs=directories["outputs"],
        temporary=directories["tmp"],
        runtime_bin=directories["runtime"],
        skill_roots=(),
        limits=SandboxLimits(runtime="native", python_executable=Path(sys.executable)),
    )

    assert launch.engine == "bubblewrap"
    assert not any(
        launch.argv[index : index + 3] == ("--ro-bind", "/", "/")
        for index in range(len(launch.argv) - 2)
    )
    assert (
        "--ro-bind",
        str(directories["inputs"].resolve()),
        str(directories["inputs"].resolve()),
    ) in tuple(
        launch.argv[index : index + 3]
        for index in range(len(launch.argv) - 2)
    )


def test_windows_auto_retains_container_sandbox_backend(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.cowork.sandbox.sys.platform", "win32")
    monkeypatch.setattr(
        "app.cowork.sandbox.shutil.which",
        lambda name: "C:/Program Files/Docker/docker.exe" if name == "docker" else None,
    )
    directories = {
        name: tmp_path / name for name in ("inputs", "work", "outputs", "tmp", "runtime")
    }
    for directory in directories.values():
        directory.mkdir()

    launch = build_sandbox_launch(
        command="$WORKPILOT_PYTHON build.py",
        inputs=directories["inputs"],
        work=directories["work"],
        outputs=directories["outputs"],
        temporary=directories["tmp"],
        runtime_bin=directories["runtime"],
        skill_roots=(),
        limits=SandboxLimits(
            runtime="auto",
            memory_mb=256,
            pids_limit=64,
            cpus=0.5,
        ),
    )

    assert launch.engine == "docker"
    assert "--memory" in launch.argv and "256m" in launch.argv
    assert "--pids-limit" in launch.argv and "64" in launch.argv
    assert "--cpus" in launch.argv and "0.5" in launch.argv


def test_office_edit_candidate_is_bound_to_source_path_and_baseline(tmp_path: Path) -> None:
    outputs = tmp_path / "outputs"
    workspace = tmp_path / "workspace"
    outputs.mkdir()
    workspace.mkdir()
    source = workspace / "report.docx"
    original = Document()
    original.add_heading("原始文档", level=1)
    original.add_paragraph("baseline")
    original.save(source)
    baseline = _sha256(source)
    candidate = Document()
    candidate.add_heading("编辑后的文档", level=1)
    candidate.add_paragraph("updated")
    candidate.save(outputs / "report.docx")
    unrelated = Document()
    unrelated.add_heading("未声明的新文件", level=1)
    unrelated.save(outputs / "new.docx")

    committed, warnings = _commit_outputs(
        outputs,
        workspace,
        max_files=10,
        max_bytes=10 * 1024 * 1024,
        office_edit_baselines={"report.docx": baseline},
    )

    assert committed == (str(source),)
    assert _sha256(source) != baseline
    assert any("候选路径必须与声明的既有源文件完全一致" in warning for warning in warnings)

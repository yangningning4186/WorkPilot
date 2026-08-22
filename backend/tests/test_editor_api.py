from collections.abc import AsyncIterator
from pathlib import Path

import httpx
import pytest
from docx import Document

from app.api.dependencies import (
    get_editor_permission_store,
    get_model_gateway,
    require_owner_identity,
)
from app.core.config import Settings, get_settings
from app.core.db import DbSession as AsyncSession
from app.core.db import get_db_session
from app.main import create_app
from app.rag.editor_permissions import EditorPermissionStore
from tests.fakes import DeterministicProvider
from workpilot_ai.gateway import ModelGateway

pytestmark = pytest.mark.integration


class MemoryEditorPermissionStore(EditorPermissionStore):
    def __init__(self) -> None:
        self.tokens: set[str] = set()

    async def grant(self, token: str, *, ttl_s: int) -> None:
        assert ttl_s > 0
        self.tokens.add(token)

    async def ttl(self, token: str) -> int:
        return 3600 if token in self.tokens else -2

    async def revoke(self, token: str) -> None:
        self.tokens.discard(token)


async def test_workspace_write_requires_grant_then_executes_directly(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    library = tmp_path / "library"
    library.mkdir()
    document = Document()
    document.add_paragraph("旧内容")
    document.save(str(library / "brief.docx"))
    settings = Settings.model_validate(
        {
            "cowork_default_workspace_path": library,
            "admin_cookie_name": "workpilot_admin_session",
        }
    )
    store = MemoryEditorPermissionStore()
    gateway = ModelGateway(
        DeterministicProvider(
            completion_text=(
                '{"summary":"更新内容","operations":['
                '{"op":"replace_paragraph","paragraph":0,"text":"新内容"}]}'
            )
        ),
        embedding_dimensions=1024,
    )

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[require_owner_identity] = lambda: None
    app.dependency_overrides[get_editor_permission_store] = lambda: store
    app.dependency_overrides[get_model_gateway] = lambda: gateway
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        client.cookies.set("workpilot_admin_session", "owner-token")
        files = (await client.get("/api/v1/editor/files")).json()["items"]
        file_id = files[0]["file_id"]
        current = (await client.get(f"/api/v1/editor/files/{file_id}")).json()
        body = {
            "baseline_sha256": current["baseline_sha256"],
            "instruction": "把第一段改成新内容",
        }

        denied = await client.post(f"/api/v1/editor/files/{file_id}/execute", json=body)
        granted = await client.post("/api/v1/editor/permission")
        applied = await client.post(f"/api/v1/editor/files/{file_id}/execute", json=body)
        revoked = await client.delete("/api/v1/editor/permission")

    assert denied.status_code == 403
    assert granted.status_code == 200
    assert granted.json()["scope"] == "local_office_write"
    assert applied.status_code == 200
    assert applied.json()["change_count"] == 1
    assert revoked.status_code == 204
    assert Document(str(library / "brief.docx")).paragraphs[0].text == "新内容"

import zipfile
from pathlib import Path
from uuid import UUID

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook  # type: ignore[import-untyped]
from openpyxl.chart import BarChart, Reference  # type: ignore[import-untyped]
from openpyxl.styles import Font  # type: ignore[import-untyped]
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings
from app.cowork.office_workspace import (
    DocumentNotEditableError,
    OfficePlanError,
    WorkspaceFileTooLargeError,
    _ensure_excel_roundtrip_safe,
    _read_raw_snapshot,
    _scan_cowork_office_root,
    _validate_formula,
    execute_cowork_office_instruction,
    execute_workspace_instruction,
    get_cowork_office_file,
    get_workspace_file,
    list_workspace_files,
)
from app.cowork.permissions import CapabilityDeniedError, create_session_root
from app.rag.local_dir import register_local_dir
from app.runstore.runs import ensure_conversation
from tests.fakes import DeterministicProvider
from workpilot_ai.gateway import ModelGateway

pytestmark = pytest.mark.integration


def test_excel_formula_rejects_network_and_external_workbook_references() -> None:
    with pytest.raises(OfficePlanError):
        _validate_formula('=HYPERLINK("https://example.com","open")')
    with pytest.raises(OfficePlanError):
        _validate_formula("='[other.xlsx]Sheet1'!A1")
    with pytest.raises(OfficePlanError, match="DDE"):
        _validate_formula("=cmd|'/C calc'!A0")
    with pytest.raises(OfficePlanError, match="安全白名单"):
        _validate_formula("=SHELL(A1)")
    _validate_formula("=SUM(A1:A3)+ROUND(B1,2)")


def test_excel_with_chart_is_rejected_before_openpyxl_roundtrip(tmp_path: Path) -> None:
    path = tmp_path / "chart.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["月份", "金额"])
    sheet.append(["一月", 10])
    chart = BarChart()
    chart.add_data(Reference(sheet, min_col=2, min_row=1, max_row=2), titles_from_data=True)
    sheet.add_chart(chart, "D2")
    workbook.save(path)

    with pytest.raises(DocumentNotEditableError, match="图表"):
        _ensure_excel_roundtrip_safe(path)


@pytest.mark.parametrize(
    ("part", "label"),
    [
        ("xl/externalLinks/externalLink1.xml", "外部链接"),
        ("xl/tables/table1.xml", "Excel 表"),
        ("xl/threadedComments/threadedComment1.xml", "线程批注"),
    ],
)
def test_excel_with_unpreserved_parts_is_rejected(
    tmp_path: Path,
    part: str,
    label: str,
) -> None:
    path = tmp_path / "unsafe.xlsx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(part, "<xml />")

    with pytest.raises(DocumentNotEditableError, match=label):
        _ensure_excel_roundtrip_safe(path)


def test_snapshot_rejects_oversized_file_before_reading_it(tmp_path: Path) -> None:
    path = tmp_path / "large.xlsx"
    with path.open("wb") as stream:
        stream.truncate(2 * 1024 * 1024)

    with pytest.raises(WorkspaceFileTooLargeError):
        _read_raw_snapshot(path, 1024)


def test_cowork_scan_is_bounded_and_skips_dependencies_and_backups(tmp_path: Path) -> None:
    root = tmp_path / "workspace"
    root.mkdir()
    (root / "brief.docx").write_bytes(b"docx")
    (root / "node_modules").mkdir()
    (root / "node_modules" / "ignored.xlsx").write_bytes(b"xlsx")
    backup = root / ".workpilot-backups" / "generation"
    backup.mkdir(parents=True)
    (backup / "ignored.docx").write_bytes(b"docx")

    items = _scan_cowork_office_root(
        UUID(int=1), "workspace", root, max_files=10, max_scan_entries=10
    )

    assert [item.relative_path for item in items] == ["brief.docx"]


async def _office_fixture(
    session: AsyncSession, tmp_path: Path
) -> tuple[Path, Settings, dict[str, str]]:
    library = tmp_path / "library"
    library.mkdir()
    (library / "note.md").write_text("# Note\n\n需要整理", encoding="utf-8")

    word = Document()
    paragraph = word.add_paragraph()
    run = paragraph.add_run("旧标题")
    run.bold = True
    word.add_paragraph("保留段落")
    table = word.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "旧单元格"
    word.save(str(library / "brief.docx"))

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "预算"
    sheet["A1"] = "项目"
    sheet["A1"].font = Font(bold=True)
    sheet["A2"] = "模型"
    sheet["B2"] = 10
    workbook.save(library / "budget.xlsx")

    settings = Settings.model_validate({"local_library_path": library})
    await register_local_dir(
        session, requested_root=Path("."), allowed_root=library, name="办公资料"
    )
    items = await list_workspace_files(session, settings=settings)
    return library, settings, {item.kind: item.file_id for item in items}


async def test_workspace_lists_and_previews_markdown_word_and_excel(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    _, settings, file_ids = await _office_fixture(db_session, tmp_path)

    assert set(file_ids) == {"markdown", "word", "excel"}
    word = await get_workspace_file(db_session, file_id=file_ids["word"], settings=settings)
    excel = await get_workspace_file(db_session, file_id=file_ids["excel"], settings=settings)

    assert "[段落 0 · 样式 Normal] 旧标题" in word.content
    assert "[表 0 · 0,0] 旧单元格" in word.content
    assert "[预算!A1] 项目" in excel.content
    assert "[预算!B2] 10" in excel.content


async def test_word_instruction_writes_directly_preserves_main_run_style_and_backs_up(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    library, settings, file_ids = await _office_fixture(db_session, tmp_path)
    word_path = library / "brief.docx"
    word_path.chmod(0o640)
    current = await get_workspace_file(
        db_session, file_id=file_ids["word"], settings=settings
    )
    provider = DeterministicProvider(
        completion_text=(
            '{"summary":"更新标题和表格","operations":['
            '{"op":"replace_paragraph","paragraph":0,"text":"新标题"},'
            '{"op":"replace_table_cell","table":0,"row":0,"column":0,"text":"新单元格"}'
            "]}"
        )
    )
    gateway = ModelGateway(provider, embedding_dimensions=1024)

    result = await execute_workspace_instruction(
        db_session,
        gateway,
        file_id=file_ids["word"],
        baseline_sha256=current.baseline_sha256,
        instruction="更新标题和表格",
        content=None,
        selection_start=0,
        selection_end=0,
        settings=settings,
    )

    changed = Document(str(word_path))
    assert changed.paragraphs[0].text == "新标题"
    assert changed.paragraphs[0].runs[0].bold is True
    assert changed.tables[0].cell(0, 0).text == "新单元格"
    assert result.change_count == 2
    assert result.backup_uri is not None
    assert (library / result.backup_uri).is_file()
    assert word_path.stat().st_mode & 0o777 == 0o640


async def test_backups_keep_only_recent_versions_per_file(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    library, base_settings, file_ids = await _office_fixture(db_session, tmp_path)
    settings = base_settings.model_copy(update={"workspace_backup_versions_per_file": 2})

    for index in range(3):
        current = await get_workspace_file(
            db_session, file_id=file_ids["word"], settings=settings
        )
        gateway = ModelGateway(
            DeterministicProvider(
                completion_text=(
                    '{"summary":"更新标题","operations":['
                    f'{{"op":"replace_paragraph","paragraph":0,"text":"标题 {index}"}}]}}'
                )
            ),
            embedding_dimensions=1024,
        )
        await execute_workspace_instruction(
            db_session,
            gateway,
            file_id=file_ids["word"],
            baseline_sha256=current.baseline_sha256,
            instruction=f"更新为标题 {index}",
            content=None,
            selection_start=0,
            selection_end=0,
            settings=settings,
        )

    backups = list((library / ".workpilot-backups").glob("*/brief.docx"))
    assert len(backups) == 2


@pytest.mark.parametrize(
    ("kind", "valid_plan", "invalid_plan"),
    [
        (
            "word",
            '{"summary":"有效修改","operations":['
            '{"op":"replace_paragraph","paragraph":0,"text":"已修改"}]}',
            '{"summary":"无效修改","operations":['
            '{"op":"replace_paragraph","paragraph":999,"text":"越界"}]}',
        ),
        (
            "excel",
            '{"summary":"有效修改","operations":['
            '{"op":"set_cell","sheet":"预算","cell":"B2","value":20}]}',
            '{"summary":"无效修改","operations":['
            '{"op":"set_cell","sheet":"不存在","cell":"B2","value":30}]}',
        ),
    ],
)
async def test_invalid_office_plan_does_not_consume_backup_slot(
    kind: str,
    valid_plan: str,
    invalid_plan: str,
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    library, base_settings, file_ids = await _office_fixture(db_session, tmp_path)
    settings = base_settings.model_copy(update={"workspace_backup_versions_per_file": 1})
    current = await get_workspace_file(db_session, file_id=file_ids[kind], settings=settings)
    await execute_workspace_instruction(
        db_session,
        ModelGateway(DeterministicProvider(completion_text=valid_plan), embedding_dimensions=1024),
        file_id=file_ids[kind],
        baseline_sha256=current.baseline_sha256,
        instruction="先执行一次有效修改",
        content=None,
        selection_start=0,
        selection_end=0,
        settings=settings,
    )
    source = library / ("brief.docx" if kind == "word" else "budget.xlsx")
    backups_before = {
        path: path.read_bytes()
        for path in (library / ".workpilot-backups").glob(f"*/{source.name}")
    }
    assert len(backups_before) == 1
    source_before = source.read_bytes()
    latest = await get_workspace_file(db_session, file_id=file_ids[kind], settings=settings)

    with pytest.raises(OfficePlanError):
        await execute_workspace_instruction(
            db_session,
            ModelGateway(
                DeterministicProvider(completion_text=invalid_plan), embedding_dimensions=1024
            ),
            file_id=file_ids[kind],
            baseline_sha256=latest.baseline_sha256,
            instruction="执行无效修改",
            content=None,
            selection_start=0,
            selection_end=0,
            settings=settings,
        )

    assert source.read_bytes() == source_before
    assert {
        path: path.read_bytes()
        for path in (library / ".workpilot-backups").glob(f"*/{source.name}")
    } == backups_before


async def test_excel_instruction_updates_values_and_formulas_without_touching_style(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    library, settings, file_ids = await _office_fixture(db_session, tmp_path)
    workbook_path = library / "budget.xlsx"
    workbook_path.chmod(0o640)
    current = await get_workspace_file(
        db_session, file_id=file_ids["excel"], settings=settings
    )
    provider = DeterministicProvider(
        completion_text=(
            '{"summary":"补齐预算计算","operations":['
            '{"op":"set_cell","sheet":"预算","cell":"B2","value":20},'
            '{"op":"set_cell","sheet":"预算","cell":"C2","value":"=B2*2"}'
            "]}"
        )
    )
    gateway = ModelGateway(provider, embedding_dimensions=1024)

    result = await execute_workspace_instruction(
        db_session,
        gateway,
        file_id=file_ids["excel"],
        baseline_sha256=current.baseline_sha256,
        instruction="把模型预算改为 20，并在 C2 计算两倍",
        content=None,
        selection_start=0,
        selection_end=0,
        settings=settings,
    )

    workbook = load_workbook(workbook_path, data_only=False)
    try:
        assert workbook["预算"]["B2"].value == 20
        assert workbook["预算"]["C2"].value == "=B2*2"
        assert workbook["预算"]["A1"].font.bold is True
    finally:
        workbook.close()
    assert result.change_count == 2
    assert result.backup_uri is not None
    assert workbook_path.stat().st_mode & 0o777 == 0o640


async def test_cowork_office_entry_requires_session_root_capability_then_writes_directly(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    library, settings, _ = await _office_fixture(db_session, tmp_path)
    word_path = library / "brief.docx"
    conversation_id = await ensure_conversation(
        db_session, scope="local_owner", title="Office capability"
    )
    await db_session.commit()

    with pytest.raises(CapabilityDeniedError):
        await get_cowork_office_file(
            db_session,
            conversation_id=conversation_id,
            target_path=word_path,
            settings=settings,
        )

    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(library),
        access_mode="read_only",
    )
    await db_session.commit()
    current = await get_cowork_office_file(
        db_session,
        conversation_id=conversation_id,
        target_path=word_path,
        settings=settings,
    )
    gateway = ModelGateway(
        DeterministicProvider(
            completion_text=(
                '{"summary":"更新标题","operations":['
                '{"op":"replace_paragraph","paragraph":0,"text":"Cowork 新标题"}]}'
            )
        ),
        embedding_dimensions=1024,
    )
    with pytest.raises(CapabilityDeniedError):
        await execute_cowork_office_instruction(
            db_session,
            gateway,
            conversation_id=conversation_id,
            target_path=word_path,
            baseline_sha256=current.baseline_sha256,
            instruction="更新标题",
            kind="word",
            settings=settings,
        )

    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(library),
        access_mode="read_write",
    )
    await db_session.commit()
    applied, authorization = await execute_cowork_office_instruction(
        db_session,
        gateway,
        conversation_id=conversation_id,
        target_path=word_path,
        baseline_sha256=current.baseline_sha256,
        instruction="更新标题",
        kind="word",
        settings=settings,
    )

    assert applied.change_count == 1
    assert authorization.capability == "office.word.edit"
    assert Document(str(word_path)).paragraphs[0].text == "Cowork 新标题"

    excel_path = library / "budget.xlsx"
    excel_current = await get_cowork_office_file(
        db_session,
        conversation_id=conversation_id,
        target_path=excel_path,
        settings=settings,
    )
    excel_gateway = ModelGateway(
        DeterministicProvider(
            completion_text=(
                '{"summary":"更新预算","operations":['
                '{"op":"set_cell","sheet":"预算","cell":"B2","value":88}]}'
            )
        ),
        embedding_dimensions=1024,
    )
    excel_applied, excel_authorization = await execute_cowork_office_instruction(
        db_session,
        excel_gateway,
        conversation_id=conversation_id,
        target_path=excel_path,
        baseline_sha256=excel_current.baseline_sha256,
        instruction="把预算改为 88",
        kind="excel",
        settings=settings,
    )
    changed_workbook = load_workbook(excel_path, data_only=False)
    try:
        assert changed_workbook["预算"]["B2"].value == 88
    finally:
        changed_workbook.close()
    assert excel_applied.change_count == 1
    assert excel_authorization.capability == "office.excel.edit"

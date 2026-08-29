import asyncio
import json
import shlex
import sys
import time
from collections.abc import AsyncIterator
from datetime import UTC, datetime, timedelta
from pathlib import Path
from uuid import UUID

import httpx
import pytest
from docx import Document
from pydantic import BaseModel, ConfigDict
from sqlalchemy.ext.asyncio import AsyncEngine
from uuid6 import uuid7

from app.agent_core.session_records import reduce_model_step_attempts
from app.api.dependencies import (
    get_run_bus,
    get_run_queue_dependency,
    require_owner_identity,
)
from app.core.config import get_settings
from app.core.db import DbSession as AsyncSession
from app.core.db import SessionFactory, get_db_session, session_factory
from app.core.run_bus import InMemoryRunBus
from app.cowork.authorization import arguments_sha256
from app.cowork.automation_tools import register_scheduler_tools
from app.cowork.browser_tools import register_browser_tools
from app.cowork.context_usage import get_cowork_context_usage
from app.cowork.interactions import enqueue_queued_message, enqueue_steering
from app.cowork.memory import load_visible_memories, remember
from app.cowork.memory_tools import register_memory_tools
from app.cowork.permissions import (
    CapabilityDeniedError,
    create_session_root,
    grant_capability,
    list_session_roots,
)
from app.cowork.runtime import (
    CoworkCheckpointCorruptionError,
    CoworkHookBus,
    ModelAttemptHookContext,
    _completion_record_payload,
    _encode_tool_result,
    _external_action_sha256,
    _independent_board_assignment_batch,
    initialize_cowork_state,
    load_cowork_checkpoint,
)
from app.cowork.schedules import claim_due_sleeping_runs, list_schedules
from app.cowork.semantic_approvals import (
    SEMANTIC_REVIEW_DENIAL_MESSAGE,
    build_trusted_approval_evidence,
)
from app.cowork.tools import (
    CoworkToolContext,
    CoworkToolError,
    CoworkToolRegistry,
    CoworkToolResult,
    CoworkToolSpec,
    RunShellArgs,
    _trusted_artifact_mime_type,
    build_default_cowork_registry,
)
from app.cowork_store.routing import cowork_store
from app.main import create_app
from app.rag.kb import local_kb_service
from app.runstore.conversations import get_conversation, update_conversation_runtime
from app.runstore.runs import (
    append_message,
    create_run,
    ensure_conversation,
    get_run,
    list_events,
    reap_expired_runs,
    request_cancel,
)
from app.worker.cowork_run import _cowork_error_detail, _cowork_failure_message, cowork_run
from app.worker.maintenance import next_run_dispatch_tick
from tests.conftest import iso_ago
from tests.fakes import DeterministicProvider
from workpilot_ai.gateway import ModelGateway
from workpilot_ai.providers.openai_compatible import (
    OpenAICompatibleProvider,
    ProviderContextOverflowError,
    ProviderTimeoutError,
)
from workpilot_ai.routing import Tier
from workpilot_ai.types import (
    CompletionChunk,
    CompletionResult,
    Message,
    ToolCall,
    ToolDefinition,
    Usage,
)

pytestmark = pytest.mark.integration


def test_cowork_failure_message_keeps_actionable_detail_bounded() -> None:
    assert _cowork_failure_message("最新 checkpoint 不是 Cowork v2 state") == (
        "Cowork 执行失败：最新 checkpoint 不是 Cowork v2 state"
    )
    assert len(_cowork_failure_message("x" * 1000)) < 380
    assert _cowork_error_detail(TimeoutError()) == "模型或工具请求超时，请重试"


def test_external_approval_action_hash_is_stable_and_non_null() -> None:
    first = _external_action_sha256("publish", {"b": 2, "a": 1})
    second = _external_action_sha256("publish", {"a": 1, "b": 2})

    assert first == second
    assert len(first) == 64


def test_artifact_mime_must_match_trusted_extension() -> None:
    assert _trusted_artifact_mime_type(Path("report.xml"), None) == "application/xml"
    with pytest.raises(CoworkToolError, match="必须与扩展名一致"):
        _trusted_artifact_mime_type(Path("payload.xml"), "text/html")


def test_office_goal_uses_skill_and_shell_without_format_specific_tools() -> None:
    registry = build_default_cowork_registry()
    names = {
        definition.name for definition in registry.tool_definitions_for("帮我生成一个儿童节 PPT")
    }

    assert {"write_file", "run_shell"} <= names
    assert {
        "create_native_artifact",
        "list_office_files",
        "inspect_office_file",
        "edit_word",
        "edit_excel",
    }.isdisjoint(names)


def latest_checkpoint_state(store_sql, run_id) -> dict:
    """SQLite 里 checkpoint 的 state 是 TEXT，取出来要自己反序列化。"""

    rows = store_sql(
        """SELECT state FROM agent_checkpoints
           WHERE run_id = ? ORDER BY checkpoint_id DESC LIMIT 1""",
        (str(run_id),),
    )
    return json.loads(rows[0]["state"])


async def test_new_cowork_run_inherits_assistant_and_tool_history(
    db_session: AsyncSession,
    tmp_path: Path,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(
        db_session,
        title="Cowork 跨 run 上下文",
    )
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    first = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="搜索今天的 AI 热点",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=first.goal,
        run_id=first.id,
    )
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=first.id, registry=registry)
    checkpoint = await load_cowork_checkpoint(db_session, run_id=first.id)
    assert checkpoint is not None
    checkpoint.state["messages"].extend(
        [
            {
                "role": "assistant",
                "content": "",
                "tool_calls": [
                    {
                        "id": "web-1",
                        "type": "function",
                        "function": {
                            "name": "web_search",
                            "arguments": '{"query":"AI 热点"}',
                        },
                    }
                ],
            },
            {
                "role": "tool",
                "tool_call_id": "web-1",
                "content": '{"ok":true,"result":{"items":["新闻一"]}}',
            },
            {"role": "assistant", "content": "这里是今天的五条 AI 热点。"},
        ]
    )
    checkpoint.state["status"] = "done"
    store_sql(
        """UPDATE agent_checkpoints SET state = ?
           WHERE run_id = ? AND checkpoint_id = ?""",
        (
            json.dumps(checkpoint.state, ensure_ascii=False),
            str(first.id),
            checkpoint.checkpoint_id,
        ),
    )
    store_sql("UPDATE agent_runs SET status = 'done' WHERE id = ?", (str(first.id),))
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="assistant",
        content="这里是今天的五条 AI 热点。",
        run_id=first.id,
    )

    second = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="把上面的新闻总结成一份文档",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=second.goal,
        run_id=second.id,
    )
    state = await initialize_cowork_state(
        db_session,
        run_id=second.id,
        registry=registry,
    )

    assert [message["role"] for message in state["messages"]] == [
        "user",
        "assistant",
        "tool",
        "assistant",
        "user",
    ]
    assert state["messages"][2]["tool_call_id"] == "web-1"
    assert state["messages"][3]["content"] == "这里是今天的五条 AI 热点。"
    assert state["messages"][-1]["content"] == second.goal

    usage = await get_cowork_context_usage(
        db_session,
        conversation_id=conversation_id,
        settings=get_settings(),
        rag=local_kb_service(get_settings()),
    )
    assert usage["used_tokens"] > 0
    assert usage["trigger_ratio"] == 0.85
    assert usage["trigger_tokens"] < usage["context_window_tokens"]
    assert usage["breakdown"]["tool_activity"] > 0


async def test_empty_conversation_context_usage_accounts_for_lazy_tool_manifest(
    db_session: AsyncSession,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork empty context")

    usage = await get_cowork_context_usage(
        db_session,
        conversation_id=conversation_id,
        settings=get_settings(),
        rag=local_kb_service(get_settings()),
    )

    breakdown = usage["breakdown"]
    assert set(breakdown) == {
        "system",
        "tool_manifest",
        "tools",
        "loaded_tools",
        "messages",
        "tool_activity",
    }
    assert breakdown["tool_manifest"] > 0
    assert breakdown["tools"] > 0
    assert breakdown["loaded_tools"] == 0
    assert breakdown["tool_activity"] == 0
    assert sum(breakdown.values()) == usage["used_tokens"]


def test_cowork_tool_result_structurally_truncates_content_and_keeps_baseline() -> None:
    baseline = "a" * 64
    result = CoworkToolResult(
        content={
            "file_id": "opaque-file-id",
            "content": "正文" * 10_000,
            "baseline_sha256": baseline,
            "editable": True,
        }
    )
    encoded = _encode_tool_result(result, 1_000)
    payload = json.loads(encoded)

    assert len(encoded) <= 1_000
    assert payload["result"]["baseline_sha256"] == baseline
    assert payload["result"]["content_truncated"] is True
    assert payload["result"]["content_original_chars"] == 20_000
    assert len(payload["result"]["content"]) < 20_000


def test_cowork_tool_result_truncation_keeps_whole_lines_and_line_cap() -> None:
    content = "".join(f"line-{index:04d}\n" for index in range(2_500))

    payload = json.loads(
        _encode_tool_result(CoworkToolResult(content={"content": content}), 50_000)
    )

    truncated = payload["result"]["content"]
    assert payload["result"]["content_truncated"] is True
    assert payload["result"]["content_original_lines"] == 2_500
    assert len(truncated.splitlines()) == 2_000
    assert truncated.endswith("\n")
    assert truncated.splitlines()[-1] == "line-1999"


def test_cowork_tool_result_byte_limit_never_returns_a_partial_line() -> None:
    content = "".join(f"第{index:04d}行-完整内容\n" for index in range(300))

    encoded = _encode_tool_result(CoworkToolResult(content={"content": content}), 1_000)
    payload = json.loads(encoded)
    truncated = payload["result"]["content"]

    assert len(encoded) <= 1_000
    assert len(encoded.encode("utf-8")) <= 1_000
    assert truncated.endswith("\n")
    assert all(line.endswith("-完整内容") for line in truncated.splitlines())


def test_failed_structured_result_keeps_output_but_marks_protocol_error() -> None:
    encoded = _encode_tool_result(
        CoworkToolResult(content={"exit_code": 1, "stdout": "", "stderr": "bad"}),
        1_000,
        result_error="Shell 命令退出码 1",
    )
    payload = json.loads(encoded)

    assert payload["ok"] is False
    assert payload["error"] == "Shell 命令退出码 1"
    assert payload["result"]["exit_code"] == 1
    assert payload["result"]["stderr"] == "bad"


def test_tool_result_separates_model_content_from_runtime_details() -> None:
    result = CoworkToolResult(
        content={"summary": "读取 3 行"},
        details={
            "row_count": 3,
            "panel": "table",
            "internal_cursor": "cursor-secret",
        },
    )

    payload = json.loads(_encode_tool_result(result, 1_000))

    assert payload["result"] == {"summary": "读取 3 行"}
    assert "cursor-secret" not in json.dumps(payload, ensure_ascii=False)
    stored = result.stored()
    assert "output" not in stored
    assert stored["content"] == {"summary": "读取 3 行"}
    assert stored["details"] == {
        "row_count": 3,
        "panel": "table",
        "internal_cursor": "cursor-secret",
    }


async def test_shell_rejects_read_only_cwd(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Read-only shell cwd")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_only",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="查看当前目录",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    context = CoworkToolContext(
        session=db_session,
        gateway=ModelGateway(DeterministicProvider(), embedding_dimensions=1024),
        settings=get_settings().model_copy(update={"cowork_shell_allowlist": ["pwd"]}),
        conversation_id=conversation_id,
        run_id=run.id,
        worker_id="readonly-shell-worker",
        plan_step_id=UUID(int=43),
        tool_call_id="readonly-shell-call",
    )

    with pytest.raises(CapabilityDeniedError, match=r"filesystem\.write"):
        await build_default_cowork_registry().execute(
            "run_shell",
            {"command": "pwd", "cwd": str(tmp_path), "reason": "查看目录"},
            context=context,
        )


async def test_browser_open_requires_scoped_network_on_top_of_browser_read(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    """browser_open 的返回值就是完整页面快照，不能绕过 origin scope。"""

    conversation_id = await ensure_conversation(db_session, title="Browser capability split")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="browser.read",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="打开网页",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    registry = build_default_cowork_registry()
    register_browser_tools(registry)
    context = CoworkToolContext(
        session=db_session,
        gateway=ModelGateway(DeterministicProvider(), embedding_dimensions=1024),
        settings=get_settings(),
        conversation_id=conversation_id,
        run_id=run.id,
        worker_id="browser-capability-worker",
        plan_step_id=UUID(int=71),
        tool_call_id="browser-capability-call",
    )

    # 校验发生在 handler 之前，所以这里不会真的去启动 Chromium。
    with pytest.raises(CapabilityDeniedError, match=r"network\.fetch"):
        await registry.execute(
            "browser_open",
            {"url": "https://example.com/"},
            context=context,
        )


async def test_shell_requires_executor_approval_and_reuses_completed_invocation(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Shell executor guard")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="host.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="执行有副作用的诊断命令",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await db_session.commit()

    output = tmp_path / "shell-executions.txt"
    script = (
        "from pathlib import Path; "
        f"p=Path({str(output)!r}); p.open('a', encoding='utf-8').write('run\\n')"
    )
    command = f"{shlex.quote(sys.executable)} -c {shlex.quote(script)}"
    arguments = {"command": command, "cwd": str(tmp_path), "reason": "验证 Shell 幂等边界"}
    registry = build_default_cowork_registry()
    gateway = ModelGateway(DeterministicProvider(), embedding_dimensions=1024)
    settings = get_settings().model_copy(update={"cowork_shell_allowlist": []})
    base_context = dict(
        session=db_session,
        gateway=gateway,
        settings=settings,
        conversation_id=conversation_id,
        run_id=run.id,
        worker_id="shell-guard-worker",
        plan_step_id=UUID(int=42),
        tool_call_id="shell-guard-call",
        semantic_approval_signing_key="3" * 64,
        cancel_event=None,
    )

    with pytest.raises(CoworkToolError, match="缺少可验证的本次审批证据"):
        await registry.execute(
            "run_shell",
            arguments,
            context=CoworkToolContext(**base_context),
        )
    assert not output.exists()

    canonical_arguments = RunShellArgs.model_validate(arguments).model_dump(mode="json")
    approved_context = CoworkToolContext(
        **base_context,
        approved_call_ids=frozenset({"shell-guard-call"}),
        approval_evidence={
            "shell-guard-call": build_trusted_approval_evidence(
                signing_key="3" * 64,
                source="user",
                run_id=run.id,
                tool_call_id="shell-guard-call",
                tool="run_shell",
                arguments_sha256=arguments_sha256(canonical_arguments),
                details={"inbox_id": str(UUID(int=43)), "standing_rule_id": None},
            )
        },
    )
    with pytest.raises(CoworkToolError, match="参数在批准后发生变化"):
        await registry.execute(
            "run_shell",
            {**arguments, "reason": "批准后被替换的参数"},
            context=approved_context,
        )
    first = await registry.execute("run_shell", arguments, context=approved_context)
    replay = await registry.execute("run_shell", arguments, context=approved_context)

    assert first.reused is False
    assert replay.reused is True
    assert first.authorization_receipt is not None
    assert replay.authorization_receipt is not None
    assert first.authorization_receipt["approval"]["source"] == "user"
    assert first.authorization_receipt["decisions"][0]["grant_id"] is not None
    assert output.read_text(encoding="utf-8") == "run\n"


async def test_shell_long_output_keeps_tail_and_registers_full_output_artifact(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Shell full output artifact")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="host.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行长输出命令",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    registry = build_default_cowork_registry()
    payload_parts = ["x" * 20] * 150
    payload = "".join(payload_parts)
    command = "/usr/bin/printf %s " + " ".join(map(shlex.quote, payload_parts))
    result = await registry.execute(
        "run_shell",
        {
            "command": command,
            "cwd": str(tmp_path),
            "reason": "验证长输出保留策略",
        },
        context=CoworkToolContext(
            session=db_session,
            gateway=ModelGateway(DeterministicProvider(), embedding_dimensions=1024),
            settings=get_settings().model_copy(
                update={
                    "cowork_shell_allowlist": ["/usr/bin/printf"],
                    "cowork_shell_max_output_bytes": 1_024,
                    "cowork_shell_full_output_max_bytes": 10_000,
                }
            ),
            conversation_id=conversation_id,
            run_id=run.id,
            worker_id="shell-full-output-worker",
            plan_step_id=UUID(int=44),
            tool_call_id="shell-full-output-call",
        ),
    )

    assert result.output["output_truncated"] is True
    assert result.output["stdout"] == payload[-1_024:]
    full_output_path = Path(result.output["full_output_path"])
    assert await asyncio.to_thread(full_output_path.is_file)
    full_output_text = await asyncio.to_thread(full_output_path.read_text, encoding="utf-8")
    assert payload in full_output_text
    assert result.output["full_output_artifact_id"]
    full_artifact = next(
        item
        for item in result.output["artifacts"]
        if item["artifact_id"] == result.output["full_output_artifact_id"]
    )
    assert full_artifact["file"]["path"] == str(full_output_path)


class RecordingCoworkQueue:
    def __init__(self) -> None:
        self.run_ids: list[UUID] = []
        self.attempts: list[int] = []

    async def enqueue_cowork_run(self, run_id: UUID, *, attempt: int = 0) -> None:
        self.run_ids.append(run_id)
        self.attempts.append(attempt)


def _tool_completion(*calls: ToolCall) -> CompletionResult:
    return CompletionResult(
        text="",
        model="fake-chat",
        provider="deterministic_test",
        usage=Usage(input_tokens=3, output_tokens=2),
        tool_calls=tuple(calls),
    )


def _final_completion(text: str) -> CompletionResult:
    return CompletionResult(
        text=text,
        model="fake-chat",
        provider="deterministic_test",
        usage=Usage(input_tokens=3, output_tokens=2),
    )


class NativeToolProvider(DeterministicProvider):
    def __init__(
        self,
        tool_completions: list[CompletionResult],
        *,
        regular_completions: list[str] | None = None,
    ) -> None:
        super().__init__(completion_texts=regular_completions)
        self.tool_completions = list(tool_completions)
        self.tool_calls = 0
        self.tool_histories: list[list[Message]] = []
        self.last_tools: list[ToolDefinition] = []
        self.parallel_flags: list[bool] = []
        self.regular_calls = 0
        self.regular_histories: list[list[Message]] = []

    async def complete(
        self,
        messages: list[Message],
        *,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        self.regular_calls += 1
        self.regular_histories.append(messages)
        return await super().complete(
            messages,
            max_tokens=max_tokens,
            temperature=temperature,
        )

    async def complete_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        del max_tokens, temperature
        self.tool_calls += 1
        self.tool_histories.append(messages)
        self.last_tools = tools
        self.parallel_flags.append(parallel_tool_calls)
        return self.tool_completions.pop(0)


class CitationEscalationGateway(ModelGateway):
    """Small routed gateway double that exposes which tier produced each candidate."""

    def __init__(
        self,
        *,
        main: list[CompletionResult],
        heavy: list[CompletionResult],
    ) -> None:
        super().__init__(DeterministicProvider(), embedding_dimensions=1024)
        self.completions = {"main": list(main), "heavy": list(heavy)}
        self.tier_calls: list[Tier] = []
        self.histories: list[list[Message]] = []

    def escalation_plan(self, task_type: str) -> tuple[Tier | None, Tier | None]:
        assert task_type == "cowork_decision"
        return ("main", "heavy")

    async def complete_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool,
        task_type: str = "generate",
        max_tokens: int = 1024,
        temperature: float = 0.0,
        tier_override: Tier | None = None,
    ) -> CompletionResult:
        del tools, parallel_tool_calls, task_type, max_tokens, temperature
        tier: Tier = tier_override or "main"
        self.tier_calls.append(tier)
        self.histories.append(messages)
        return self.completions[tier].pop(0)

    async def stream_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool = True,
        task_type: str = "generate",
        max_tokens: int = 1024,
        temperature: float = 0.0,
        tier_override: Tier | None = None,
    ) -> AsyncIterator[CompletionChunk]:
        result = await self.complete_with_tools(
            messages,
            tools=tools,
            parallel_tool_calls=parallel_tool_calls,
            task_type=task_type,
            max_tokens=max_tokens,
            temperature=temperature,
            tier_override=tier_override,
        )
        yield CompletionChunk(result=result)


async def test_truncated_text_is_retried_and_never_delivered(
    db_engine: AsyncEngine,
    db_session: AsyncSession,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Truncated model text")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="给我一份完整答复",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            CompletionResult(
                text="这是绝不能交付的半截答复",
                model="fake-chat",
                provider="deterministic_test",
                usage=Usage(input_tokens=3, output_tokens=2),
                stop_reason="length",
            ),
            _final_completion("这是重新生成的完整答复。"),
        ]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "done"
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    assert checkpoint.state["final_message"] == "这是重新生成的完整答复。"
    assert provider.tool_calls == 2
    assert any(
        message.role == "user"
        and "model_output_truncated" in message.content
        and "重新发送一份完整决策" in message.content
        for message in provider.tool_histories[1]
    )
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert any(
        event.type == "error" and event.payload.get("code") == "model_output_truncated"
        for event in events
    )


async def test_truncated_tool_batch_is_denied_without_executing_any_call(
    db_engine: AsyncEngine,
    db_session: AsyncSession,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Truncated tool batch")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="维护任务清单",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    truncated = _tool_completion(
        ToolCall(
            id="truncated-todo",
            name="todo_write",
            arguments=json.dumps(
                {"todos": [{"content": "不应执行", "status": "in_progress"}]},
                ensure_ascii=False,
            ),
        )
    )
    truncated = CompletionResult(
        text=truncated.text,
        model=truncated.model,
        provider=truncated.provider,
        usage=truncated.usage,
        tool_calls=truncated.tool_calls,
        stop_reason="length",
    )
    provider = NativeToolProvider([truncated, _final_completion("已安全重试。")])

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    assert checkpoint.state["todos"] == []
    assert checkpoint.state["model_truncation_retries"] == 0
    denial = next(
        message
        for message in checkpoint.state["messages"]
        if message.get("role") == "tool" and message.get("tool_call_id") == "truncated-todo"
    )
    assert "本批调用均未执行" in denial["content"]
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert not any(
        event.type == "tool.start" and event.payload.get("tool") == "todo_write" for event in events
    )


async def test_failed_citation_repair_escalates_main_candidate_to_heavy(
    db_engine: AsyncEngine,
    db_session: AsyncSession,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Citation escalation")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="根据知识库解释这个结论",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        kb_slug="test-kb",
    )
    state = latest_checkpoint_state(store_sql, run.id)
    state["evidence_ledger"] = [
        {
            "ledger_id": "ledger-k1",
            "citation_id": "K1",
            "kind": "knowledge",
            "block_id": "block-k1",
            "version_id": "version-k1",
            "document_id": "document-k1",
            "title": "测试资料",
            "source_uri": "file:///test.md",
            "quote": "这是经过读取的原文证据。",
            "quote_sha256": "a" * 64,
            "char_start": 0,
            "char_end": 12,
            "heading_path": ["结论"],
            "locations": [],
            "material_id": None,
            "locator": None,
            "verified": True,
            "tool_call_id": "seed-evidence",
        }
    ]
    store_sql(
        "UPDATE agent_checkpoints SET state = ? WHERE run_id = ?",
        (json.dumps(state, ensure_ascii=False), str(run.id)),
    )
    gateway = CitationEscalationGateway(
        main=[
            _final_completion("第一份草稿没有提供任何可回查的知识引用。"),
            _final_completion("修复后的主档草稿仍然漏掉了知识引用。"),
        ],
        heavy=[_final_completion("重档重新生成了可交付的结论。[K1]")],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": gateway,
            "cowork_registry": registry,
        },
        str(run.id),
    )

    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    assert checkpoint.state["status"] == "done"
    assert checkpoint.state["final_message"].endswith("[K1]")
    assert gateway.tier_calls == ["main", "main", "heavy"]
    assert checkpoint.state["citation_repair_attempts"] == 1
    assert checkpoint.state["runtime_snapshot"]["model_identities"] == [
        "deterministic_test/fake-chat"
    ]


class SemanticReviewOutcomeProvider(NativeToolProvider):
    def __init__(
        self,
        tool_completions: list[CompletionResult],
        review_outcomes: list[str | Exception],
    ) -> None:
        super().__init__(tool_completions)
        self.review_outcomes = list(review_outcomes)

    async def complete(
        self,
        messages: list[Message],
        *,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        del max_tokens, temperature
        self.regular_calls += 1
        self.regular_histories.append(messages)
        outcome = self.review_outcomes.pop(0)
        if isinstance(outcome, Exception):
            raise outcome
        return CompletionResult(
            text=outcome,
            model="semantic-review-test",
            provider="deterministic_test",
            usage=Usage(input_tokens=3, output_tokens=2),
        )


async def test_title_generation_cannot_block_cowork_terminal(
    db_session: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="新会话")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="hello",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)

    async def never_finishes(*args, **kwargs) -> str:
        del args, kwargs
        await asyncio.Event().wait()
        return "不会生成的标题"

    monkeypatch.setattr("app.worker.cowork_run.generate_conversation_title", never_finishes)
    provider = NativeToolProvider([_final_completion("Hello! 有什么我可以帮你的吗？")])

    await asyncio.wait_for(
        cowork_run(
            {
                "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
                "session_factory": session_factory,
                "bus": bus,
                "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
                "cowork_registry": registry,
            },
            str(run.id),
        ),
        timeout=1,
    )

    finished = await get_run(db_session, run.id)
    assert finished is not None and finished.status == "done"
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert [event.type for event in events][-4:] == [
        "message.snapshot",
        "message.done",
        "agent.end",
        "run.done",
    ]
    terminal_progress = [
        event
        for event in events
        if event.type == "step.update" and event.payload.get("status") == "done"
    ]
    assert terminal_progress[-1].payload.get("summary") == ""


async def test_durable_follow_up_continues_the_same_run_after_a_final_answer(
    db_session: AsyncSession,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Durable follow-up")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="先回答第一问",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    queued = await enqueue_queued_message(
        db_session,
        run_id=run.id,
        conversation_id=conversation_id,
        content="这是同一轮的追问",
        source="local_owner",
        delivery="follow_up",
    )
    provider = NativeToolProvider(
        [_final_completion("第一问答案"), _final_completion("追问答案")]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    finished = await get_run(db_session, run.id)
    assert finished is not None and finished.status == "done"
    assert provider.tool_calls == 2
    second_history = provider.tool_histories[1]
    assert any(message.role == "assistant" and message.content == "第一问答案" for message in second_history)
    assert any(
        message.role == "user" and message.content == "这是同一轮的追问"
        for message in second_history
    )
    rows = await cowork_store().export_rows(
        table="cowork_steering_messages",
        columns=("id", "delivery", "status"),
    )
    persisted = next(row for row in rows if row["id"] == str(queued.id))
    assert persisted == {
        "id": str(queued.id),
        "delivery": "follow_up",
        "status": "consumed",
    }
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert any(event.type == "queue.message.applied" for event in events)
    action_records = [
        record
        for record in await cowork_store().list_session_records(run_id=run.id)
        if record.kind == "harness_action"
    ]
    assert [record.payload["action"] for record in action_records if record.phase == "completed"] == [
        "prepare",
        "dispatch",
        "materialize",
        "prepare",
        "dispatch",
        "materialize",
    ]
    assert all(record.phase != "failed" for record in action_records)


async def test_durable_next_run_dispatch_is_idempotent_and_preserves_provenance(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Durable next run")
    terminal = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="已完成的上一轮",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    assert await cowork_store().finish_run(run_id=terminal.id, status="done")
    queued = await enqueue_queued_message(
        db_session,
        run_id=terminal.id,
        conversation_id=conversation_id,
        content="从终态启动下一轮",
        source="local_owner",
        delivery="follow_up",
    )
    assert queued.delivery == "next_run" and queued.status == "ready"
    queue = RecordingCoworkQueue()
    bus = InMemoryRunBus()
    context = {
        "settings": get_settings().model_copy(
            update={"cowork_default_workspace_path": tmp_path / "workspace"}
        ),
        "session_factory": session_factory,
        "run_queue": queue,
        "bus": bus,
    }

    assert await next_run_dispatch_tick(context) == 1
    assert await next_run_dispatch_tick(context) == 0
    assert len(queue.run_ids) == 1
    successor = await cowork_store().get_run(queue.run_ids[0])
    assert successor is not None
    assert successor.goal == queued.content
    rows = await cowork_store().export_rows(
        table="agent_runs",
        columns=("id", "source_wake_id"),
    )
    successor_row = next(row for row in rows if row["id"] == str(successor.id))
    assert successor_row["source_wake_id"] == str(queued.id)
    checkpoint = await load_cowork_checkpoint(db_session, run_id=successor.id)
    assert checkpoint is not None
    assert checkpoint.state["semantic_review_user_text_source"] == "local_owner"
    assert checkpoint.state["messages"][-1]["source"] == "local_owner"


class TimeoutThenResumeProvider(NativeToolProvider):
    def __init__(self, outcomes: list[CompletionResult | Exception]) -> None:
        super().__init__([])
        self.outcomes = list(outcomes)

    async def complete_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        self.tool_calls += 1
        self.tool_histories.append(messages)
        self.last_tools = tools
        self.parallel_flags.append(parallel_tool_calls)
        outcome = self.outcomes.pop(0)
        if isinstance(outcome, Exception):
            raise outcome
        return outcome


async def test_route_timeout_retries_from_checkpoint_without_replaying_completed_tool(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork timeout recovery")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="列出文件后总结",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = TimeoutThenResumeProvider(
        [
            _tool_completion(
                ToolCall(
                    id="list-before-timeout",
                    name="list_files",
                    arguments=json.dumps({"path": "."}),
                )
            ),
            ProviderTimeoutError("主模型与 fallback 均超时"),
            _final_completion("已从 checkpoint 恢复并完成。"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(
            update={
                "run_heartbeat_s": 60.0,
                "run_max_recovery": 2,
                "cowork_provider_timeout_retry_base_s": 0.01,
                "cowork_provider_timeout_retry_max_s": 0.02,
            }
        ),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }

    await cowork_run(context, str(run.id))

    parked = await get_run(db_session, run.id)
    assert parked is not None and parked.status == "sleeping"
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    assert checkpoint.state["messages"][-1]["tool_call_id"] == "list-before-timeout"
    first_events = await list_events(db_session, run_id=run.id, limit=200)
    assert sum(event.type == "tool.result" for event in first_events) == 1
    retry_event = next(
        event
        for event in first_events
        if event.type == "step.update" and event.payload.get("status") == "recovering"
    )
    assert "checkpoint" in retry_event.payload["summary"]

    woken = await claim_due_sleeping_runs(
        db_session,
        now=datetime.now(UTC) + timedelta(seconds=1),
    )
    assert woken == [run.id]
    await cowork_run(context, str(run.id))

    finished = await get_run(db_session, run.id)
    assert finished is not None and finished.status == "done"
    final_history = provider.tool_histories[-1]
    assert any(message.tool_call_id == "list-before-timeout" for message in final_history)
    final_events = await list_events(db_session, run_id=run.id, limit=200)
    assert sum(event.type == "tool.result" for event in final_events) == 1


class CancellingCoworkProvider(NativeToolProvider):
    """在第一次模型决策返回后模拟用户点击停止。"""

    def __init__(
        self,
        session_factory: SessionFactory,
        run_id: UUID,
    ) -> None:
        super().__init__(
            [
                _tool_completion(
                    ToolCall(
                        id="cancel-call",
                        name="list_files",
                        arguments=json.dumps({"path": "."}),
                    )
                )
            ]
        )
        self.session_factory = session_factory
        self.run_id = run_id

    async def complete_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        result = await super().complete_with_tools(
            messages,
            tools=tools,
            parallel_tool_calls=parallel_tool_calls,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        if self.tool_calls == 1:
            async with self.session_factory() as session:
                await request_cancel(session, run_id=self.run_id)
                await session.commit()
        return result


class OverflowRecoveringCoworkProvider(NativeToolProvider):
    def __init__(
        self,
        tool_completions: list[CompletionResult],
        *,
        overflow_calls: set[int],
        regular_completions: list[str],
    ) -> None:
        super().__init__(tool_completions, regular_completions=regular_completions)
        self.overflow_calls = overflow_calls

    async def complete_with_tools(
        self,
        messages: list[Message],
        *,
        tools: list[ToolDefinition],
        parallel_tool_calls: bool,
        max_tokens: int,
        temperature: float,
    ) -> CompletionResult:
        del max_tokens, temperature
        self.tool_calls += 1
        self.tool_histories.append(messages)
        self.last_tools = tools
        self.parallel_flags.append(parallel_tool_calls)
        if self.tool_calls in self.overflow_calls:
            raise ProviderContextOverflowError("模型服务返回 400：context_length_exceeded")
        return self.tool_completions.pop(0)


class EmptyToolArgs(BaseModel):
    model_config = ConfigDict(extra="forbid")


class _FakeBrowserControl:
    def __init__(self) -> None:
        self.clicked = False

    async def click(self, **options: int) -> None:
        assert options["timeout"] > 0
        self.clicked = True


class _FakeBrowserBody:
    async def inner_text(self, **options: int) -> str:
        assert options["timeout"] > 0
        return "测试页面"


class _FakeBrowserControls:
    async def count(self) -> int:
        return 0


class _FakeBrowserPage:
    url = "https://example.com/after-click"

    def locator(self, selector: str) -> object:
        return _FakeBrowserBody() if selector == "body" else _FakeBrowserControls()

    async def wait_for_load_state(self, state: str, **options: int) -> None:
        assert state == "domcontentloaded"
        assert options["timeout"] > 0

    async def title(self) -> str:
        return "测试页面"

    async def screenshot(self, *, path: str, full_page: bool) -> None:
        assert isinstance(full_page, bool)
        await asyncio.to_thread(Path(path).write_bytes, b"\x89PNG\r\n\x1a\nworkpilot-test-image")


class _FakeBrowserSession:
    def __init__(self, control: _FakeBrowserControl) -> None:
        self.page = _FakeBrowserPage()
        self.controls: list[object] = [control]
        self.action_no = 0
        self.last_used = 0.0


class _FakeBrowserManager:
    def __init__(self, session_id: str, session: _FakeBrowserSession) -> None:
        self.session_id = session_id
        self.session = session

    async def get(
        self,
        session_id: str,
        *,
        conversation_id: UUID,
    ) -> _FakeBrowserSession:
        assert session_id == self.session_id
        assert isinstance(conversation_id, UUID)
        return self.session


async def test_browser_screenshot_reaches_next_model_turn_as_image_attachment(
    db_engine: AsyncEngine,
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Browser screenshot attachment")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="browser.read",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="截图并检查页面",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    session_id = "test-browser-screenshot"
    screenshot_path = tmp_path / "page.png"
    registry = build_default_cowork_registry()
    register_browser_tools(
        registry,
        _FakeBrowserManager(session_id, _FakeBrowserSession(_FakeBrowserControl())),  # type: ignore[arg-type]
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="load-browser-screenshot",
                    name="load_tools",
                    arguments=json.dumps({"names": ["browser_screenshot"]}),
                )
            ),
            _tool_completion(
                ToolCall(
                    id="browser-screenshot",
                    name="browser_screenshot",
                    arguments=json.dumps(
                        {
                            "session_id": session_id,
                            "path": str(screenshot_path),
                            "full_page": True,
                        }
                    ),
                )
            ),
            _final_completion("已检查截图。"),
        ]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    attachment_messages = [message for message in provider.tool_histories[2] if message.attachments]
    assert len(attachment_messages) == 1
    attachment = attachment_messages[0].attachments[0]
    assert attachment.kind == "image"
    assert attachment.path == str(screenshot_path.resolve())
    assert attachment.media_type == "image/png"
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    directive = next(
        message
        for message in checkpoint.state["messages"]
        if message.get("role") == "runtime_directive"
        and message.get("source") == "tool_result_attachment"
    )
    assert directive["attachments"][0]["sha256"] == attachment.sha256


async def test_auto_mode_browser_destructive_clicks_without_inbox_round_trip(
    db_engine: AsyncEngine, db_session: AsyncSession
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork browser control")
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="browser.destructive",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="点击浏览器控件",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    session_id = "test-browser-session-01"
    control = _FakeBrowserControl()
    registry = build_default_cowork_registry()
    register_browser_tools(
        registry,
        _FakeBrowserManager(session_id, _FakeBrowserSession(control)),  # type: ignore[arg-type]
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="load-browser-click",
                    name="load_tools",
                    arguments=json.dumps({"names": ["browser_click"]}),
                )
            ),
            _tool_completion(
                ToolCall(
                    id="browser-click",
                    name="browser_click",
                    arguments=json.dumps({"session_id": session_id, "control_index": 0}),
                )
            ),
            _final_completion("点击完成。"),
        ],
        regular_completions=['{"decision":"allow"}'],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "done"
    assert control.clicked is True
    loaded_tail = provider.tool_histories[1][-1].content
    assert "<loaded_tools>" in loaded_tail
    assert "browser_click" in loaded_tail
    assert "fetch_url" in loaded_tail and "web_search" in loaded_tail
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert "interrupt" not in {event.type for event in events}
    semantic_review = next(event for event in events if event.type == "approval.semantic_review")
    assert semantic_review.payload["decision"] == "allow"
    assert provider.regular_calls == 1
    review_envelope = json.loads(provider.regular_histories[0][-1].content)
    assert set(review_envelope) == {
        "schema",
        "frozen_session_facts",
        "user_authored_text",
        "user_text_truncated",
        "canonical_action",
    }
    assert review_envelope["canonical_action"]["target"] is not None

    follow_up = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="继续操作当前页面",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=follow_up.goal,
        run_id=follow_up.id,
    )
    resumed_registry = build_default_cowork_registry()
    register_browser_tools(resumed_registry)
    resumed = await initialize_cowork_state(
        db_session,
        run_id=follow_up.id,
        registry=resumed_registry,
        bus=bus,
    )
    assert "browser_click" in resumed["runtime_snapshot"]["tool_registry"]["activated_tools"]


async def test_batched_exclusive_browser_actions_are_rejected_without_execution(
    db_engine: AsyncEngine, db_session: AsyncSession
) -> None:
    conversation_id = await ensure_conversation(
        db_session, title="Cowork exclusive browser control"
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="browser.destructive",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="连续点击两个浏览器控件",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    session_id = "test-browser-session-02"
    control = _FakeBrowserControl()
    registry = build_default_cowork_registry()
    register_browser_tools(
        registry,
        _FakeBrowserManager(session_id, _FakeBrowserSession(control)),  # type: ignore[arg-type]
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="load-browser-click",
                    name="load_tools",
                    arguments=json.dumps({"names": ["browser_click"]}),
                )
            ),
            _tool_completion(
                ToolCall(
                    id="browser-click-1",
                    name="browser_click",
                    arguments=json.dumps({"session_id": session_id, "control_index": 0}),
                ),
                ToolCall(
                    id="browser-click-2",
                    name="browser_click",
                    arguments=json.dumps({"session_id": session_id, "control_index": 0}),
                ),
            ),
            _final_completion("已改为逐个操作。"),
        ]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    assert control.clicked is False
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    tool_errors = [
        json.loads(message["content"])
        for message in checkpoint.state["messages"]
        if message["role"] == "tool" and '"ok":false' in message["content"]
    ]
    assert len(tool_errors) == 2
    assert all("需要审批的外部动作必须单独调用" in item["error"] for item in tool_errors)
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert "interrupt" not in {event.type for event in events}
    assert not any(
        event.type == "tool.start" and event.payload.get("tool") == "browser_click"
        for event in events
    )


async def test_cowork_runner_edits_docx_via_shell_and_registers_artifact(
    db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    workspace = tmp_path / "cowork"
    workspace.mkdir()
    source_path = workspace / "brief.docx"
    output_path = workspace / "brief-workpilot.docx"
    document = Document()
    document.add_paragraph("旧内容")
    document.save(source_path)
    (workspace / "edit_docx.py").write_text(
        """from docx import Document

source = Document("brief.docx")
source.paragraphs[0].text = "新内容"
source.save("brief-workpilot.docx")
verified = Document("brief-workpilot.docx")
assert verified.paragraphs[0].text == "新内容"
print("verified: brief-workpilot.docx")
""",
        encoding="utf-8",
    )

    conversation_id = await ensure_conversation(db_session, title="Cowork Shell Office")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(workspace),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="把 brief.docx 第一段改成新内容，结果另存",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    command = f"{shlex.quote(sys.executable)} edit_docx.py"
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="shell-office",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": command,
                            "cwd": str(workspace),
                            "reason": "运行已写入工作区的 DOCX 编辑与验证脚本",
                        }
                    ),
                )
            ),
            _final_completion("已生成并验证 brief-workpilot.docx。"),
        ]
    )
    settings = get_settings().model_copy(
        update={
            "cowork_shell_allowlist": [shlex.quote(sys.executable)],
            "run_heartbeat_s": 60.0,
        }
    )

    await cowork_run(
        {
            "settings": settings,
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    assert Document(output_path).paragraphs[0].text == "新内容"
    assert Document(source_path).paragraphs[0].text == "旧内容"
    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "done"
    events = await list_events(db_session, run_id=run.id, limit=200)
    event_types = [event.type for event in events]
    assert event_types.count("tool.start") == 1
    assert event_types.count("tool.result") == 1
    assert "artifact" in event_types
    artifact = store_sql("SELECT uri, meta FROM artifacts WHERE run_id = ?", (str(run.id),))[0]
    assert artifact["uri"] == str(output_path)
    assert json.loads(artifact["meta"])["discovered_after"] == "run_shell"
    invocation = store_sql(
        """SELECT status, effect_ref FROM tool_invocations
           WHERE run_id = ? AND tool_name = 'run_shell'""",
        (str(run.id),),
    )[0]
    assert invocation["status"] == "succeeded"
    assert invocation["effect_ref"].startswith("shell:")
    shell_result = json.loads(
        next(
            message.content
            for message in provider.tool_histories[-1]
            if message.role == "tool" and message.tool_call_id == "shell-office"
        )
    )
    assert shell_result["result"]["artifact_scan"]["registered"] == 1
    assert shell_result["result"]["artifacts"][0]["file"]["path"] == str(output_path)


async def test_cowork_runner_consumes_cancel_before_pending_tool(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql, message_status
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork cancel")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="扫描文件，但我可能随时停止",
        budget_tokens=50_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = CancellingCoworkProvider(session_factory, run.id)
    settings = get_settings().model_copy(update={"run_heartbeat_s": 60.0})
    await cowork_run(
        {
            "settings": settings,
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "cancelled"
    assert provider.tool_calls == 1
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert "tool.start" not in [event.type for event in events]
    assert any(
        event.type == "step.update" and event.payload.get("status") == "skipped" for event in events
    )
    assert any(
        event.type == "error" and event.payload.get("code") == "cancelled" for event in events
    )
    assert any(
        event.type == "run.done" and event.payload.get("status") == "cancelled" for event in events
    )
    assistant_status = (await message_status(run.conversation_id, run.id, "assistant"))[0]
    assert assistant_status == "cancelled"
    latest_state = latest_checkpoint_state(store_sql, run.id)
    assert latest_state["messages"][-2]["tool_calls"][0]["id"] == "cancel-call"
    assert latest_state["messages"][-1]["role"] == "tool"
    assert latest_state["messages"][-1]["tool_call_id"] == "cancel-call"


async def test_cowork_runner_executes_parallel_safe_read_batch_concurrently(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    active = 0
    max_active = 0
    both_started = asyncio.Event()

    async def parallel_read(context: CoworkToolContext, raw: BaseModel) -> CoworkToolResult:
        nonlocal active, max_active
        del context, raw
        active += 1
        max_active = max(max_active, active)
        if active == 2:
            both_started.set()
        try:
            await asyncio.wait_for(both_started.wait(), timeout=1)
            await asyncio.sleep(0.02)
            return CoworkToolResult(content={"active": active})
        finally:
            active -= 1

    registry = CoworkToolRegistry()
    for name in ("read_a", "read_b"):
        registry.register(
            CoworkToolSpec(
                name=name,
                description=f"并行读取 {name}",
                args_model=EmptyToolArgs,
                risk="read",
                effect="none",
                parallel_safe=True,
                handler=parallel_read,
            )
        )

    conversation_id = await ensure_conversation(db_session, title="Cowork parallel reads")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="并行读取两个来源",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(id="read-a-call", name="read_a", arguments="{}"),
                ToolCall(id="read-b-call", name="read_b", arguments="{}"),
            ),
            _final_completion("两个来源均已读取。"),
        ]
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    "cowork_decision_max_tokens": 2048,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "done"
    assert max_active == 2
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert [event.type for event in events].count("tool.start") == 2
    assert [event.type for event in events].count("tool.result") == 2
    latest_state = latest_checkpoint_state(store_sql, run.id)
    canonical = [message for message in latest_state["messages"] if message["role"] != "user"]
    assert canonical[0]["role"] == "assistant"
    assert [call["id"] for call in canonical[0]["tool_calls"]] == [
        "read-a-call",
        "read-b-call",
    ]
    assert [message["tool_call_id"] for message in canonical[1:3]] == [
        "read-a-call",
        "read-b-call",
    ]
    assert canonical[3]["role"] == "assistant"
    assert canonical[3]["content"] == "两个来源均已读取。"
    assert canonical[3]["stop_reason"] == "stop"
    assert canonical[3]["created_at"]


async def test_cowork_runner_executes_independent_board_assignments_concurrently(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    del db_engine

    class AssignmentArgs(BaseModel):
        model_config = ConfigDict(extra="forbid")

        task_id: str
        worker: str

    active = 0
    max_active = 0
    both_started = asyncio.Event()

    async def assign(context: CoworkToolContext, raw: BaseModel) -> CoworkToolResult:
        nonlocal active, max_active
        del context
        args = AssignmentArgs.model_validate(raw.model_dump())
        active += 1
        max_active = max(max_active, active)
        if active == 2:
            both_started.set()
        try:
            await asyncio.wait_for(both_started.wait(), timeout=1)
            return CoworkToolResult(content={"task_id": args.task_id, "worker": args.worker})
        finally:
            active -= 1

    registry = CoworkToolRegistry()
    registry.register(
        CoworkToolSpec(
            name="board_assign_task",
            description="分配独立团队任务",
            args_model=AssignmentArgs,
            risk="write",
            effect="store",
            parallel_safe=False,
            handler=assign,
        )
    )
    conversation_id = await ensure_conversation(db_session, title="Concurrent team assignments")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="并发分配两个独立 Worker",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    first_task = str(uuid7())
    second_task = str(uuid7())
    calls = (
        ToolCall(
            id="assign-architecture",
            name="board_assign_task",
            arguments=json.dumps({"task_id": first_task, "worker": "architecture"}),
        ),
        ToolCall(
            id="assign-testing",
            name="board_assign_task",
            arguments=json.dumps({"task_id": second_task, "worker": "testing"}),
        ),
    )
    assert _independent_board_assignment_batch(
        [
            {
                "call_id": call.id,
                "name": call.name,
                "arguments": call.arguments,
                "step_idx": index,
                "step_id": str(uuid7()),
            }
            for index, call in enumerate(calls)
        ]
    )
    provider = NativeToolProvider(
        [_tool_completion(*calls), _final_completion("两个 Worker 均已返回。")]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    completed = await get_run(db_session, run.id)
    assert completed is not None and completed.status == "done"
    assert registry.get("board_assign_task").parallel_safe is False
    assert max_active == 2


async def test_cowork_runner_has_no_tool_step_count_limit(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    class NumberedReadArgs(BaseModel):
        model_config = ConfigDict(extra="forbid")

        index: int

    async def numbered_read(context: CoworkToolContext, raw: BaseModel) -> CoworkToolResult:
        del context
        args = NumberedReadArgs.model_validate(raw.model_dump())
        return CoworkToolResult(content={"index": args.index})

    registry = CoworkToolRegistry()
    registry.register(
        CoworkToolSpec(
            name="numbered_read",
            description="按序读取一个测试项",
            args_model=NumberedReadArgs,
            risk="read",
            effect="none",
            parallel_safe=True,
            handler=numbered_read,
        )
    )
    conversation_id = await ensure_conversation(db_session, title="Cowork unlimited tool steps")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="读取三十一个互不重复的测试项",
        budget_tokens=100_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    calls = tuple(
        ToolCall(
            id=f"numbered-read-{index}",
            name="numbered_read",
            arguments=json.dumps({"index": index}),
        )
        for index in range(31)
    )
    provider = NativeToolProvider([_tool_completion(*calls), _final_completion("全部读取完成。")])

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "done"
    state = latest_checkpoint_state(store_sql, run.id)
    assert state["iteration"] == 31
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert [event.type for event in events].count("tool.result") == 31


async def test_cowork_todo_list_lands_in_state_and_is_pinned_for_the_next_turn(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    """清单要进 checkpoint 并钉进下一轮 system prompt，而不是只留在历史里。"""

    conversation_id = await ensure_conversation(db_session, title="Cowork todo list")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="分三步整理资料",
        budget_tokens=60_000,
        budget_calls=20,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)

    def _todo_call(call_id: str, first_status: str) -> ToolCall:
        return ToolCall(
            id=call_id,
            name="todo_write",
            arguments=json.dumps(
                {
                    "todos": [
                        {"content": "读取源文件", "status": first_status},
                        {"content": "生成报告", "status": "pending"},
                    ]
                },
                ensure_ascii=False,
            ),
        )

    provider = NativeToolProvider(
        [
            _tool_completion(_todo_call("todo-1", "in_progress")),
            _tool_completion(_todo_call("todo-2", "completed")),
            _final_completion("两步都已完成。"),
        ]
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    "cowork_decision_max_tokens": 2048,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "done"

    latest_state = latest_checkpoint_state(store_sql, run.id)
    # 整份替换而不是追加，且 completed 在入口就被归一成 done。
    assert latest_state["todos"] == [
        {"content": "读取源文件", "status": "done"},
        {"content": "生成报告", "status": "pending"},
    ]

    events = await list_events(db_session, run_id=run.id, limit=200)
    updates = [event for event in events if event.type == "todo.update"]
    assert [item.payload["done"] for item in updates] == [0, 1]
    assert updates[-1].payload["todos"] == latest_state["todos"]

    # 第一轮还没有清单；第二轮起它出现在 outbound 末尾的临时块里，压缩不会把它冲掉。
    assert "<current_todos>" not in provider.tool_histories[0][-1].content
    second_tail = provider.tool_histories[1][-1]
    assert second_tail.role == "user"
    assert "[>] 读取源文件" in second_tail.content
    assert "[ ] 生成报告" in second_tail.content
    assert "[x] 读取源文件" in provider.tool_histories[2][-1].content
    # 清单每轮都在变，所以它必须留在 system prompt 之外：整个 run 里 system 逐字不变，
    # provider 的前缀缓存才有意义。
    systems = {history[0].content for history in provider.tool_histories}
    assert len(systems) == 1
    assert "<current_todos>" not in systems.pop()


async def test_cowork_memory_snapshot_holds_for_a_run_and_refreshes_on_the_next_one(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    """记忆在 run 起始快照一次，run 内不再变，下一条消息才刷新。

    记忆进的是 system prompt，也就是 provider 前缀缓存的那一段。若每轮重新读库，模型
    自己调一次 remember 就会让这一轮之后的每次调用都重新计费；而且它在一次 run 里
    "知道什么"会在脚下变。代价是用户在记忆面板里的改动延后到下一条消息生效。
    """

    conversation_id = await ensure_conversation(db_session, title="Cowork memory")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    original, _ = await remember(
        db_session,
        conversation_id=conversation_id,
        scope="global",
        content="用户偏好 PDF 报告",
        key="report-format",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="记住我的偏好",
        budget_tokens=60_000,
        budget_calls=20,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    await db_session.commit()
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    register_memory_tools(registry)
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)

    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="load-remember",
                    name="load_tools",
                    arguments=json.dumps({"names": ["remember"]}),
                )
            ),
            _tool_completion(
                ToolCall(
                    id="remember-1",
                    name="remember",
                    arguments=json.dumps(
                        {
                            "content": "用户偏好 Markdown 报告",
                            "scope": "global",
                            "key": "report-format",
                        },
                        ensure_ascii=False,
                    ),
                )
            ),
            _final_completion("已经记住你偏好 Markdown 报告。"),
        ]
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    "cowork_decision_max_tokens": 2048,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "done"

    # 本 run 的每一轮拿到的都是同一份 system prompt，逐字不变。
    systems = {history[0].content for history in provider.tool_histories}
    assert len(systems) == 1
    only_system = systems.pop()
    assert "用户偏好 PDF 报告" in only_system
    assert "用户偏好 Markdown 报告" not in only_system

    events = await list_events(db_session, run_id=run.id, limit=200)
    saved = [event for event in events if event.type == "memory.saved"]
    assert len(saved) == 1
    # 覆盖同 key 记忆算 updated；事件只带版本引用，长期记忆正文不复制进事件账本。
    assert saved[0].payload["action"] == "updated"
    assert saved[0].payload["previous_memory_id"] == str(original.id)
    assert "previous_content" not in saved[0].payload
    assert "content" not in saved[0].payload["memory"]

    visible = await load_visible_memories(db_session, conversation_id=conversation_id)
    assert [item.content for item in visible] == ["用户偏好 Markdown 报告"]

    # 下一条消息（新 run）重新快照：覆盖后的新文本进来，旧的不再并存。
    followup = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="按我的偏好出一份报告",
        budget_tokens=60_000,
        budget_calls=20,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=followup.goal,
        run_id=followup.id,
    )
    await db_session.commit()
    followup_registry = build_default_cowork_registry()
    register_memory_tools(followup_registry)
    await initialize_cowork_state(
        db_session, run_id=followup.id, registry=followup_registry, bus=bus
    )
    followup_provider = NativeToolProvider([_final_completion("好的。")])
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    "cowork_decision_max_tokens": 2048,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(followup_provider, embedding_dimensions=1024),
            "cowork_registry": followup_registry,
        },
        str(followup.id),
    )

    followup_system = followup_provider.tool_histories[0][0].content
    assert "用户偏好 Markdown 报告" in followup_system
    assert "用户偏好 PDF 报告" not in followup_system


async def test_cowork_recovers_provider_context_overflow_without_mutating_canonical_history(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    for index in range(100):
        (tmp_path / f"{index:03d}-{'x' * 80}.docx").write_bytes(b"x")
    conversation_id = await ensure_conversation(db_session, title="Cowork overflow recovery")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="列出文件后总结",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = OverflowRecoveringCoworkProvider(
        [
            _tool_completion(
                ToolCall(
                    id="list-before-overflow",
                    name="list_files",
                    arguments=json.dumps({"path": str(tmp_path)}),
                )
            ),
            _final_completion("已列出文件并完成总结。"),
        ],
        overflow_calls={2},
        regular_completions=['{"summary":"已经调用 list_files 并取得文件清单。"}'],
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    # 本用例只验证 provider overflow 恢复：首次决策必须**不**触发阈值
                    # 压缩，否则强制恢复路径没有可归档的历史，测的就不是这条路了。
                    # 余量由下面 ModelGateway 的 context window 提供，不再靠压榨输出
                    # 预算来腾——那个办法有下限（ge=128），迟早会被固定前缀顶穿。
                    "cowork_decision_max_tokens": 2048,
                    "cowork_compaction_trigger_ratio": 1.0,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(
                provider,
                embedding_dimensions=1024,
                # 溢出用例靠 provider 主动抛 400 触发恢复，本地阈值不能先动手；
                # 窗口留足余量，测试就不必随固定前缀的大小反复调参。
                default_context_window_tokens=128_000,
            ),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "done"
    assert refreshed.used_calls == 4  # 决策、已发出的 400、摘要、恢复后的决策
    assert provider.tool_calls == 3
    latest_state = latest_checkpoint_state(store_sql, run.id)
    canonical = latest_state["messages"]
    assert [item["role"] for item in canonical] == ["user", "assistant", "tool", "assistant"]
    assert canonical[1]["tool_calls"][0]["id"] == "list-before-overflow"
    assert canonical[2]["tool_call_id"] == "list-before-overflow"
    assert latest_state["compaction"]["summary_upto"] == 0
    assert latest_state["compaction"]["turn_prefix_upto"] == 3
    recovery_history = provider.tool_histories[-1]
    assert any("turn_prefix_summary" in item.content for item in recovery_history)
    assert not any(item.role == "tool" for item in recovery_history)
    events = await list_events(db_session, run_id=run.id, limit=200)
    compacted = [event for event in events if event.type == "context.compacted"]
    assert len(compacted) == 1
    assert compacted[0].payload["reason"] == "provider_overflow"
    session_kinds = [
        row["kind"]
        for row in store_sql(
            "SELECT kind FROM session_entries WHERE conversation_id = ? ORDER BY seq",
            (str(conversation_id),),
        )
    ]
    assert "model_change" in session_kinds
    assert "compaction" in session_kinds
    compaction_entry = store_sql(
        "SELECT payload FROM session_entries WHERE conversation_id = ? "
        "AND kind = 'compaction' ORDER BY seq DESC LIMIT 1",
        (str(conversation_id),),
    )[0]
    compaction_payload = json.loads(compaction_entry["payload"])
    assert compaction_payload["turn_prefix_summary"]
    assert "已经调用 list_files" in compaction_payload["turn_prefix_summary"]
    assert "details" in compaction_payload
    attempts = reduce_model_step_attempts(
        await cowork_store().list_session_records(run_id=run.id)
    )
    assert any(
        attempt.step == "compaction" and attempt.phase == "completed"
        for attempt in attempts
    )


async def test_cowork_context_overflow_progress_guard_stops_recovery_loop(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    for index in range(100):
        (tmp_path / f"{index:03d}-{'x' * 80}.docx").write_bytes(b"x")
    conversation_id = await ensure_conversation(db_session, title="Cowork overflow guard")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="列出文件后继续",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = OverflowRecoveringCoworkProvider(
        [
            _tool_completion(
                ToolCall(
                    id="list-loop-guard",
                    name="list_files",
                    arguments=json.dumps({"path": str(tmp_path)}),
                )
            )
        ],
        overflow_calls={2, 3},
        regular_completions=['{"summary":"已经取得文件清单。"}'],
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={
                    # 同上：首次决策必须不触发阈值压缩，强制恢复才有历史可归档。
                    "cowork_decision_max_tokens": 2048,
                    "cowork_compaction_trigger_ratio": 1.0,
                    "cowork_context_overflow_max_recoveries": 2,
                    "run_heartbeat_s": 60.0,
                }
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(
                provider,
                embedding_dimensions=1024,
                # 溢出用例靠 provider 主动抛 400 触发恢复，本地阈值不能先动手；
                # 窗口留足余量，测试就不必随固定前缀的大小反复调参。
                default_context_window_tokens=128_000,
            ),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None
    assert refreshed.status == "failed"
    assert provider.tool_calls == 3
    events = await list_events(db_session, run_id=run.id, limit=200)
    overflow_errors = [
        event
        for event in events
        if event.type == "error" and event.payload.get("code") == "cowork_context_overflow"
    ]
    assert len(overflow_errors) == 1
    assert overflow_errors[0].payload["recoveries"] == 2


def test_default_registry_exposes_risk_and_capability_contract() -> None:
    registry = build_default_cowork_registry()
    catalog = {item["name"]: item for item in registry.catalog()}

    assert {
        "create_native_artifact",
        "list_office_files",
        "inspect_office_file",
        "edit_word",
        "edit_excel",
        "list_workspace_roots",
        "search_tool_catalog",
    }.isdisjoint(catalog)
    assert catalog["list_files"]["parallel_safe"] is True
    assert catalog["read_file"]["capability"] == "filesystem.read"
    assert catalog["search_files"]["effect"] == "none"
    assert catalog["read_file"]["parallel_safe"] is True
    assert catalog["fetch_url"]["capability"] == "network.fetch"
    assert catalog["write_file"]["effect"] == "filesystem"
    assert catalog["write_file"]["risk"] == "write"
    assert registry.parallel_safe(["list_files", "search_files", "read_file"]) is True
    assert catalog["run_shell"]["capability"] == "host.execute"
    assert catalog["run_sandbox"]["capability"] == "sandbox.execute"
    assert catalog["run_shell"]["effect"] == "external"
    assert catalog["ask_user"]["execution"] == "interaction"
    assert catalog["request_directory"]["execution"] == "interaction"
    assert catalog["request_capability"]["execution"] == "interaction"
    assert registry.is_exclusive("run_shell") is True
    assert registry.is_exclusive("sleep") is True
    assert registry.is_exclusive("ask_user") is True


@pytest.mark.parametrize(
    ("interaction_body", "expected_tool_fragment"),
    [
        ({"approved": True, "answer": "按部门"}, "按部门"),
        ({"approved": False}, "基于现有信息自行判断"),
    ],
    ids=("answered", "declined"),
)
async def test_cowork_steering_and_ask_user_pause_resume_with_canonical_tool_history(
    interaction_body: dict[str, object],
    expected_tool_fragment: str,
    db_engine: AsyncEngine,
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork interaction")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="整理材料",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)

    queue = RecordingCoworkQueue()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        steering_response = await client.post(
            f"/api/v1/runs/{run.id}/steering",
            json={"message": "先只整理财务相关材料"},
        )
    assert steering_response.status_code == 202

    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="ask-scope",
                    name="ask_user",
                    arguments=json.dumps(
                        {"question": "按月份还是按部门整理？", "choices": ["月份", "部门"]}
                    ),
                )
            ),
            _final_completion("已按部门完成整理。"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }
    await cowork_run(context, str(run.id))

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert any(
        message.role == "user" and message.content == "先只整理财务相关材料"
        for message in provider.tool_histories[0]
    )
    events = await list_events(db_session, run_id=run.id, limit=200)
    interrupt = next(event for event in events if event.type == "interrupt")
    assert interrupt.payload["kind"] == "ask_user"

    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/runs/{run.id}/interactions/{interrupt.payload['resume_token']}/respond",
            json=interaction_body,
        )
    assert response.status_code == 202
    assert response.json()["status"] == "queued"
    assert queue.run_ids == [run.id]
    assert queue.attempts[0] > 0
    if isinstance(interaction_body.get("answer"), str):
        from app.cowork_store.factory import local_cowork_stores

        records = await local_cowork_stores().conversations.read(conversation_id)
        assert not [
            item
            for item in records
            if item.role == "user"
            and str(item.run_id) == str(run.id)
            and item.content == interaction_body["answer"]
        ]

    await cowork_run(context, str(run.id))
    completed = await get_run(db_session, run.id)
    assert completed is not None and completed.status == "done"
    resumed_history = provider.tool_histories[1]
    assistant_call = next(message for message in resumed_history if message.tool_calls)
    assert assistant_call.tool_calls[0].id == "ask-scope"
    tool_result = next(
        message for message in resumed_history if message.role == "tool" and message.tool_call_id
    )
    assert tool_result.tool_call_id == "ask-scope"
    assert expected_tool_fragment in tool_result.content
    if interaction_body["approved"] is False:
        assert '"status":"rejected"' in tool_result.content


async def test_queued_message_api_exposes_effective_delivery_and_cancel(
    db_session: AsyncSession,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Queued message API")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="等待后续消息",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    bus = InMemoryRunBus()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        queued_response = await client.post(
            f"/api/v1/runs/{run.id}/queued-messages",
            json={"message": "本轮结束后继续", "delivery": "follow_up"},
        )
        assert queued_response.status_code == 202
        queued = queued_response.json()
        assert queued["requested_delivery"] == queued["delivery"] == "follow_up"
        assert queued["status"] == "pending"
        cancelled = await client.delete(
            f"/api/v1/runs/{run.id}/queued-messages/{queued['message_id']}"
        )
        assert cancelled.status_code == 204

    assert await cowork_store().finish_run(run_id=run.id, status="done")
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        successor_response = await client.post(
            f"/api/v1/runs/{run.id}/queued-messages",
            json={"message": "终态之后另起一轮", "delivery": "follow_up"},
        )
    assert successor_response.status_code == 202
    successor = successor_response.json()
    assert successor["requested_delivery"] == "follow_up"
    assert successor["delivery"] == "next_run"
    assert successor["status"] == "ready"


@pytest.mark.parametrize("tool_name", ["request_directory", "request_capability"])
async def test_cowork_runtime_permission_requests_apply_only_after_user_approval(
    tool_name: str, db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    initial_root = tmp_path / f"initial-{tool_name}"
    initial_root.mkdir()
    extra_root = tmp_path / f"extra-{tool_name}"
    extra_root.mkdir()
    conversation_id = await ensure_conversation(db_session, title=f"Cowork {tool_name}")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(initial_root),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="需要额外权限",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    arguments = (
        {"reason": "读取补充材料", "access_mode": "read_only"}
        if tool_name == "request_directory"
        else {"reason": "同步外部系统", "capability": "external.write"}
    )
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id=f"permission-{tool_name}",
                    name=tool_name,
                    arguments=json.dumps(arguments),
                )
            )
        ]
    )
    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )
    events = await list_events(db_session, run_id=run.id, limit=200)
    interrupt = next(event for event in events if event.type == "interrupt")
    assert interrupt.payload["kind"] == (
        "directory_request" if tool_name == "request_directory" else "capability_request"
    )

    queue = RecordingCoworkQueue()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    body: dict[str, object] = {"approved": True}
    if tool_name == "request_directory":
        body["path"] = str(extra_root)
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/runs/{run.id}/interactions/{interrupt.payload['resume_token']}/respond",
            json=body,
        )
    assert response.status_code == 202
    assert response.json()["status"] == "queued"
    if tool_name == "request_directory":
        granted = store_sql(
            """SELECT access_mode FROM session_roots
               WHERE conversation_id = ? AND canonical_path = ?""",
            (str(conversation_id), str(extra_root.resolve())),
        )[0]["access_mode"]
        assert granted == "read_only"
    else:
        granted = store_sql(
            """SELECT capability FROM capability_grants
               WHERE conversation_id = ?
                 AND capability = 'external.write' AND revoked_at IS NULL""",
            (str(conversation_id),),
        )[0]["capability"]
        assert granted == "external.write"


async def test_cowork_shell_defaults_missing_cwd_to_writable_root(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork shell argument recovery")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="分析工作目录",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="shell-missing-cwd",
                    name="run_shell",
                    arguments=json.dumps({"command": "pwd", "reason": "检查当前目录"}),
                )
            ),
            _final_completion("已确认当前工作目录。"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(
            update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": ["pwd"]}
        ),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }

    await cowork_run(context, str(run.id))

    completed = await get_run(db_session, run.id)
    assert completed is not None and completed.status == "done"
    assert completed.error is None
    assert provider.tool_calls == 2
    correction_history = provider.tool_histories[1]
    tool_result = next(
        message
        for message in correction_history
        if message.role == "tool" and message.tool_call_id == "shell-missing-cwd"
    )
    assert '"ok":true' in tool_result.content
    assert str(tmp_path) in tool_result.content
    events = await list_events(db_session, run_id=run.id, limit=200)
    assert any(
        event.type == "tool.result" and event.payload.get("tool") == "run_shell" for event in events
    )


async def test_cowork_nonzero_shell_exit_is_a_failed_step_with_visible_output(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork failed shell truth")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行失败命令并说明结果",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="shell-fails",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/false",
                            "reason": "验证非零退出码的状态",
                        }
                    ),
                )
            ),
            _final_completion("命令执行失败，已保留退出码供诊断。"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(
            update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": ["/usr/bin/false"]}
        ),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }

    await cowork_run(context, str(run.id))

    history = provider.tool_histories[1]
    tool_result = next(
        message
        for message in history
        if message.role == "tool" and message.tool_call_id == "shell-fails"
    )
    payload = json.loads(tool_result.content)
    assert payload["ok"] is False
    assert payload["result"]["exit_code"] == 1

    events = await list_events(db_session, run_id=run.id, limit=200)
    shell_error = next(
        event
        for event in events
        if event.type == "tool.error" and event.payload.get("tool") == "run_shell"
    )
    assert "退出码 1" in shell_error.payload["error"]
    assert shell_error.payload["activity"]["summary"] == "验证非零退出码的状态"
    assert not any(
        event.type == "tool.result" and event.payload.get("tool") == "run_shell" for event in events
    )

    attempt = store_sql(
        """SELECT status, error_model FROM agent_attempts
           WHERE run_id = ? AND tool_name = 'run_shell'""",
        (str(run.id),),
    )[0]
    assert attempt["status"] == "failed"
    assert "退出码 1" in attempt["error_model"]


async def test_cowork_office_flow_exposes_shell_and_skill_loading(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork Office shell boundary")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="整理工作空间里的 Word 文档",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider([_final_completion("需要执行时会加载 DOCX Skill 并使用 Shell。")])

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    completed = await get_run(db_session, run.id)
    assert completed is not None and completed.status == "done"
    names = {tool.name for tool in provider.last_tools}
    assert "run_shell" in names
    assert {"inspect_office_file", "edit_word", "edit_excel"}.isdisjoint(names)


async def test_cowork_shell_requires_once_approval_then_executes_exact_pending_call(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork shell approval")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行一条诊断命令",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="shell-once",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/printf shell-approved",
                            "cwd": str(tmp_path),
                            "reason": "确认本地命令执行链",
                        }
                    ),
                )
            ),
            _final_completion("诊断命令已完成。"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(
            update={
                "run_heartbeat_s": 60.0,
                "cowork_shell_allowlist": [],
            }
        ),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }
    await cowork_run(context, str(run.id))
    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    events = await list_events(db_session, run_id=run.id, limit=200)
    interrupt = next(event for event in events if event.type == "interrupt")
    assert interrupt.payload["kind"] == "shell_approval"
    assert interrupt.payload["payload"]["command"] == "/usr/bin/printf shell-approved"
    assert interrupt.payload["payload"]["allowlisted"] is False

    queue = RecordingCoworkQueue()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/runs/{run.id}/interactions/{interrupt.payload['resume_token']}/respond",
            json={"approved": True},
        )
    assert response.status_code == 202
    assert response.json()["status"] == "queued"

    await cowork_run(context, str(run.id))
    completed = await get_run(db_session, run.id)
    assert completed is not None and completed.status == "done"
    resumed_history = provider.tool_histories[1]
    tool_result = next(
        message
        for message in resumed_history
        if message.role == "tool" and message.tool_call_id == "shell-once"
    )
    assert "shell-approved" in tool_result.content
    assert '"execution_mode":"argv"' in tool_result.content


@pytest.mark.parametrize(
    "review_outcome, disposition",
    [
        ("not-json", "invalid_response"),
        (RuntimeError("provider diagnostic must not escape"), "provider_error"),
    ],
)
async def test_auto_semantic_review_failure_falls_back_to_manual_approval(
    db_session: AsyncSession,
    tmp_path: Path,
    review_outcome: str | Exception,
    disposition: str,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Semantic review fallback")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行一次诊断命令",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    provider = SemanticReviewOutcomeProvider(
        [
            _tool_completion(
                ToolCall(
                    id="review-fallback-shell",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/printf review-fallback",
                            "cwd": str(tmp_path),
                            "reason": "诊断",
                        }
                    ),
                )
            )
        ],
        [review_outcome],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": []}
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert provider.regular_calls == 1
    events = await list_events(db_session, run_id=run.id, limit=200)
    review = next(event for event in events if event.type == "approval.semantic_review")
    assert review.payload["decision"] == "unsure"
    assert review.payload["disposition"] == disposition
    assert "provider diagnostic" not in json.dumps(
        [event.payload for event in events], ensure_ascii=False
    )
    assert any(event.type == "interrupt" for event in events)


async def test_external_inbound_text_never_authorizes_auto_semantic_review(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Inbound review provenance")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="来自飞书的外部命令",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
        trace_id="messaging",
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="external_inbound",
    )
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="external-inbound-shell",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/printf inbound",
                            "cwd": str(tmp_path),
                            "reason": "外部输入",
                        }
                    ),
                )
            )
        ]
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": []}
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert provider.regular_calls == 0
    events = await list_events(db_session, run_id=run.id, limit=200)
    review = next(event for event in events if event.type == "approval.semantic_review")
    assert review.payload["decision"] == "unsure"
    assert review.payload["disposition"] == "untrusted_user_text"


async def test_local_steering_preserves_provenance_for_semantic_review(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Steering review provenance")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行诊断命令",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    # Local steering keeps provenance and is included verbatim in semantic-review evidence.
    await enqueue_steering(
        db_session,
        run_id=run.id,
        conversation_id=conversation_id,
        content="不要执行，取消刚才的命令",
        source="local_owner",
    )
    await db_session.commit()
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="post-steering-shell",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/printf stale-goal",
                            "cwd": str(tmp_path),
                            "reason": "旧目标",
                        }
                    ),
                )
            ),
            _final_completion("已取消执行。"),
        ],
        regular_completions=['{"decision":"deny"}'],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": []}
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "done"
    assert provider.regular_calls == 1
    events = await list_events(db_session, run_id=run.id, limit=200)
    review = next(event for event in events if event.type == "approval.semantic_review")
    assert review.payload["disposition"] == "reviewed"
    assert review.payload["decision"] == "deny"


async def test_truncated_owner_goal_never_reaches_semantic_reviewer(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Truncated review intent")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    # Direct/legacy callers can bypass the HTTP schema's 4k bound.  A revocation in the
    # omitted tail must not be hidden from the reviewer and then auto-allowed.
    goal = "执行诊断 " + ("x" * 4_000) + " 不要执行，取消"
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal=goal,
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="truncated-goal-shell",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": "/usr/bin/printf truncated-goal",
                            "cwd": str(tmp_path),
                            "reason": "诊断",
                        }
                    ),
                )
            )
        ],
        regular_completions=['{"decision":"allow"}'],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": []}
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert provider.regular_calls == 0
    events = await list_events(db_session, run_id=run.id, limit=200)
    review = next(event for event in events if event.type == "approval.semantic_review")
    assert review.payload["disposition"] == "truncated_user_text"


async def test_opaque_external_body_never_reaches_or_passes_semantic_reviewer(
    db_session: AsyncSession,
) -> None:
    class OpaqueExternalArgs(BaseModel):
        model_config = ConfigDict(extra="forbid")

        target: str
        body: str

    async def handler(
        context: CoworkToolContext, raw: BaseModel
    ) -> CoworkToolResult:  # pragma: no cover - approval must pause before execution
        del context, raw
        return CoworkToolResult(content={"ok": True})

    conversation_id = await ensure_conversation(db_session, title="Opaque action review")
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="向 acct-1 发送通知",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    registry.register(
        CoworkToolSpec(
            name="opaque_external_send",
            description="发送带正文的外部通知",
            args_model=OpaqueExternalArgs,
            capability=None,
            risk="external",
            effect="external",
            parallel_safe=False,
            handler=handler,
            approval_required=True,
            approval_target_fields=("target",),
        )
    )
    bus = InMemoryRunBus()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="opaque-send",
                    name="opaque_external_send",
                    arguments=json.dumps({"target": "acct-1", "body": "hidden sensitive body"}),
                )
            )
        ],
        regular_completions=['{"decision":"allow"}'],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert provider.regular_calls == 0
    events = await list_events(db_session, run_id=run.id, limit=200)
    review = next(event for event in events if event.type == "approval.semantic_review")
    assert review.payload["disposition"] == "opaque_or_truncated_action"
    assert "hidden sensitive body" not in json.dumps(review.payload, ensure_ascii=False)


async def test_consecutive_semantic_denies_trip_auto_to_manual_with_generic_agent_error(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Semantic deny breaker")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="运行诊断命令",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=registry,
        bus=bus,
        semantic_review_user_text_source="local_owner",
    )
    provider = SemanticReviewOutcomeProvider(
        [
            _tool_completion(
                ToolCall(
                    id=f"denied-shell-{index}",
                    name="run_shell",
                    arguments=json.dumps(
                        {
                            "command": f"/usr/bin/printf deny-{index}",
                            "cwd": str(tmp_path),
                            "reason": "诊断",
                        }
                    ),
                )
            )
            for index in range(1, 4)
        ],
        ['{"decision":"deny"}', '{"decision":"deny"}'],
    )

    await cowork_run(
        {
            "settings": get_settings().model_copy(
                update={"run_heartbeat_s": 60.0, "cowork_shell_allowlist": []}
            ),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert provider.regular_calls == 2
    conversation = await get_conversation(db_session, conversation_id=conversation_id)
    assert conversation is not None and conversation.approval_mode == "interactive"
    events = await list_events(db_session, run_id=run.id, limit=300)
    reviews = [event for event in events if event.type == "approval.semantic_review"]
    assert [event.payload["decision"] for event in reviews] == ["deny", "deny"]
    assert reviews[0].payload["breaker_tripped"] is False
    assert reviews[1].payload["breaker_tripped"] is True
    agent_tool_results = [
        message.content
        for history in provider.tool_histories
        for message in history
        if message.role == "tool" and message.tool_call_id in {"denied-shell-1", "denied-shell-2"}
    ]
    assert agent_tool_results
    assert all(SEMANTIC_REVIEW_DENIAL_MESSAGE in content for content in agent_tool_results)
    assert all("deny" not in content.casefold() for content in agent_tool_results)


async def test_unpersisted_semantic_breaker_cannot_fail_open_on_next_run(
    db_session: AsyncSession,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Unpersisted review breaker")
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    first = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="first",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=first.goal,
        run_id=first.id,
    )
    registry = build_default_cowork_registry()
    await initialize_cowork_state(
        db_session,
        run_id=first.id,
        registry=registry,
        semantic_review_user_text_source="local_owner",
    )
    state = latest_checkpoint_state(store_sql, first.id)
    state["semantic_review_consecutive_denies"] = 2
    state["semantic_review_breaker_tripped"] = True
    state["semantic_review_breaker_persisted"] = False
    state["status"] = "done"
    store_sql(
        "UPDATE agent_checkpoints SET state = ? WHERE run_id = ?",
        (json.dumps(state, ensure_ascii=False), str(first.id)),
    )
    store_sql("UPDATE agent_runs SET status = 'done' WHERE id = ?", (str(first.id),))

    second = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="second",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=second.goal,
        run_id=second.id,
    )
    next_state = await initialize_cowork_state(
        db_session,
        run_id=second.id,
        registry=build_default_cowork_registry(),
        semantic_review_user_text_source="local_owner",
    )

    assert next_state["semantic_review_breaker_tripped"] is True
    assert next_state["semantic_review_breaker_persisted"] is False
    assert next_state["semantic_review_consecutive_denies"] == 2


async def test_auto_mode_cannot_waive_a_protected_workspace_write(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Protected write floor")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="更新 workspace policy",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    await db_session.commit()

    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="protected-config-write",
                    name="write_file",
                    arguments=json.dumps(
                        {
                            "path": ".workpilot/config.toml",
                            "content": '[shell]\nallow = ["python"]\n',
                            "create_parents": True,
                            "purpose": "workspace",
                        }
                    ),
                )
            )
        ]
    )
    context = {
        "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }

    await cowork_run(context, str(run.id))

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert not (tmp_path / ".workpilot/config.toml").exists()
    events = await list_events(db_session, run_id=run.id, limit=200)
    interrupt = next(event for event in events if event.type == "interrupt")
    assert interrupt.payload["kind"] == "external_approval"
    assert interrupt.payload["payload"]["tool"] == "write_file"
    assert interrupt.payload["payload"]["human_only"] is True
    assert interrupt.payload["payload"]["standing_action_target"] is None
    assert provider.regular_calls == 0


async def test_auto_mode_cannot_waive_schedule_persistent_authority(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Schedule authority floor")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="external.write",
    )
    await update_conversation_runtime(
        db_session,
        conversation_id=conversation_id,
        provider_profile_id=None,
        model_override=None,
        unattended=False,
        approval_mode="auto",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="创建明天的自动化",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    await db_session.commit()

    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    register_scheduler_tools(registry)
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="load-schedule-tool",
                    name="load_tools",
                    arguments=json.dumps({"names": ["create_schedule"]}),
                )
            ),
            _tool_completion(
                ToolCall(
                    id="create-schedule-human-only",
                    name="create_schedule",
                    arguments=json.dumps(
                        {
                            "title": "明日摘要",
                            "goal": "整理项目摘要",
                            "schedule_kind": "once",
                            "run_at": "2026-09-01T09:00:00+08:00",
                            "timezone": "Asia/Shanghai",
                            "standing_approvals": [],
                        }
                    ),
                )
            ),
        ]
    )
    context = {
        "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }

    await cowork_run(context, str(run.id))

    waiting = await get_run(db_session, run.id)
    assert waiting is not None and waiting.status == "waiting_human"
    assert await list_schedules(db_session) == []
    events = await list_events(db_session, run_id=run.id, limit=200)
    interrupt = next(event for event in events if event.type == "interrupt")
    assert interrupt.payload["payload"]["tool"] == "create_schedule"
    assert interrupt.payload["payload"]["human_only"] is True
    assert provider.regular_calls == 0


async def test_cowork_cancel_terminates_running_allowlisted_shell_process(
    db_engine: AsyncEngine, db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork shell cancel")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    await grant_capability(
        db_session,
        conversation_id=conversation_id,
        capability="shell.execute",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="启动后停止诊断进程",
        budget_tokens=50_000,
        budget_calls=20,
        budget_wall_ms=120_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    bus = InMemoryRunBus()
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    sleep_script = tmp_path / "sleep-for-cancel-test.py"
    sleep_script.write_text("import time\ntime.sleep(30)\n", encoding="utf-8")
    command = f"{shlex.quote(sys.executable)} {shlex.quote(str(sleep_script))}"
    provider = NativeToolProvider(
        [
            _tool_completion(
                ToolCall(
                    id="shell-cancel",
                    name="run_shell",
                    arguments=json.dumps(
                        {"command": command, "cwd": str(tmp_path), "reason": "测试停止"}
                    ),
                )
            ),
            _final_completion("不应执行到这里"),
        ]
    )
    context = {
        "settings": get_settings().model_copy(
            update={
                "run_heartbeat_s": 60.0,
                "cowork_cancel_poll_s": 0.05,
                "cowork_shell_allowlist": [shlex.quote(sys.executable)],
                "cowork_shell_timeout_s": 60,
                "cowork_shell_terminate_grace_s": 0.1,
            }
        ),
        "session_factory": session_factory,
        "bus": bus,
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
    }
    started = time.monotonic()
    worker_task = asyncio.create_task(cowork_run(context, str(run.id)))
    deadline = time.monotonic() + 4
    while time.monotonic() < deadline:
        async with session_factory() as session:
            events = await list_events(session, run_id=run.id, limit=200)
        if any(
            event.type == "tool.start" and event.payload.get("tool") == "run_shell"
            for event in events
        ):
            break
        await asyncio.sleep(0.02)
    else:
        worker_task.cancel()
        raise AssertionError("shell 工具未在期限内启动")

    async with session_factory() as session:
        await request_cancel(session, run_id=run.id)
        await session.commit()
    await asyncio.wait_for(worker_task, timeout=5)

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "cancelled"
    assert time.monotonic() - started < 5
    assert provider.tool_calls == 1


async def test_cowork_run_api_initializes_checkpoint_and_enqueues(
    db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    conversation_id = await ensure_conversation(db_session, title="新会话")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    selected_file = tmp_path / "季度汇报.docx"
    selected_file.write_bytes(b"selected-workspace-file")
    await db_session.commit()
    queue = RecordingCoworkQueue()
    bus = InMemoryRunBus()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    settings = get_settings().model_copy(
        update={"cowork_default_workspace_path": tmp_path / "WorkPilot"}
    )
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            "/api/v1/runs/cowork",
            json={
                "conversation_id": str(conversation_id),
                "goal": "整理目录中的预算表",
                "workspace_files": [str(selected_file)],
            },
        )

    assert response.status_code == 202
    run_id = UUID(response.json()["run_id"])
    assert response.json()["workflow_type"] == "cowork"
    assert response.json()["conversation_title"] == "整理目录中的预算表"
    assert queue.run_ids == [run_id]
    checkpoint_state = latest_checkpoint_state(store_sql, run_id)
    config_rows = store_sql(
        "SELECT config FROM cowork_run_configs WHERE run_id = ?", (str(run_id),)
    )
    run_config = json.loads(config_rows[0]["config"])
    assert checkpoint_state["schema_version"] == "cowork.v3"
    assert "workspace_files" not in checkpoint_state
    assert run_config["workspace_files"] == [str(selected_file.resolve())]
    restored = await load_cowork_checkpoint(db_session, run_id=run_id)
    assert restored is not None
    assert restored.state["workspace_files"] == [str(selected_file.resolve())]


async def test_current_checkpoint_with_missing_field_is_rejected_not_repaired(
    db_session: AsyncSession,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Strict checkpoint restore")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="验证恢复边界",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=build_default_cowork_registry(),
    )
    state = latest_checkpoint_state(store_sql, run.id)
    assert state["schema_version"] == "cowork.v3"
    state.pop("model_truncation_retries")
    store_sql(
        "UPDATE agent_checkpoints SET state = ? WHERE run_id = ?",
        (json.dumps(state, ensure_ascii=False), str(run.id)),
    )

    with pytest.raises(CoworkCheckpointCorruptionError) as caught:
        await load_cowork_checkpoint(db_session, run_id=run.id)

    assert caught.value.code == "missing_fields"
    assert "model_truncation_retries" in str(caught.value)


async def test_resume_rejects_missing_model_identity_before_model_call(
    db_engine: AsyncEngine,
    db_session: AsyncSession,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Missing model identity")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="继续之前的模型会话",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    bus = InMemoryRunBus()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry, bus=bus)
    state = latest_checkpoint_state(store_sql, run.id)
    state["runtime_snapshot"]["model_identities"] = ["removed-provider/removed-model"]
    store_sql(
        "UPDATE agent_checkpoints SET state = ? WHERE run_id = ?",
        (json.dumps(state, ensure_ascii=False), str(run.id)),
    )
    provider = NativeToolProvider([_final_completion("不应被调用")])

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": bus,
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    refreshed = await get_run(db_session, run.id)
    assert refreshed is not None and refreshed.status == "failed"
    assert provider.tool_calls == 0
    events = await list_events(db_session, run_id=run.id, limit=200)
    failure = next(event for event in events if event.type == "error")
    assert "removed-provider/removed-model" in failure.payload["user_message"]


async def test_session_facts_are_inherited_after_git_remote_and_roots_change(
    db_session: AsyncSession,
    tmp_path: Path,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Frozen session facts")
    repository = tmp_path / "repository"
    git_dir = repository / ".git"
    git_dir.mkdir(parents=True)
    config = git_dir / "config"
    config.write_text(
        '[remote "origin"]\nurl = https://token@initial.example/team/private.git\n',
        encoding="utf-8",
    )
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(repository),
        access_mode="read_only",
    )
    first = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="首次运行",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    first_state = await initialize_cowork_state(
        db_session,
        run_id=first.id,
        registry=build_default_cowork_registry(),
    )
    assert first_state["session_facts"]["workspace_roots"][0]["git_remote_hostnames"] == [
        "initial.example"
    ]
    store_sql("UPDATE agent_runs SET status = 'done' WHERE id = ?", (str(first.id),))

    config.write_text(
        '[remote "origin"]\nurl = https://new-token@mutated.example/other/repo.git\n',
        encoding="utf-8",
    )
    added = tmp_path / "added-later"
    added.mkdir()
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(added),
        access_mode="read_write",
    )
    second = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="后续运行",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    second_state = await initialize_cowork_state(
        db_session,
        run_id=second.id,
        registry=build_default_cowork_registry(),
    )

    assert second_state["session_facts"] == first_state["session_facts"]
    serialized = json.dumps(second_state["session_facts"], ensure_ascii=False)
    assert "mutated.example" not in serialized
    assert str(added) not in serialized


async def test_old_checkpoint_without_session_facts_recovers_as_unavailable(
    db_session: AsyncSession,
    tmp_path: Path,
    store_sql,
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Legacy session facts")
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="恢复旧 checkpoint",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=build_default_cowork_registry(),
    )
    state = latest_checkpoint_state(store_sql, run.id)
    config_rows = store_sql(
        "SELECT config FROM cowork_run_configs WHERE run_id = ?", (str(run.id),)
    )
    state = {**json.loads(config_rows[0]["config"]), **state}
    state["schema_version"] = "cowork.v2"
    state.pop("session_facts")
    # 模拟 v22 以前的整份 checkpoint：当时没有独立 RunConfig 行。
    store_sql("DELETE FROM cowork_run_configs WHERE run_id = ?", (str(run.id),))
    store_sql(
        "UPDATE agent_checkpoints SET state = ? WHERE run_id = ?",
        (json.dumps(state, ensure_ascii=False), str(run.id)),
    )

    restored = await load_cowork_checkpoint(db_session, run_id=run.id)

    assert restored is not None
    assert restored.state["session_facts"]["capture_status"] == "legacy_unavailable"
    assert restored.state["session_facts"]["workspace_roots"] == []

    store_sql("UPDATE agent_runs SET status = 'done' WHERE id = ?", (str(run.id),))
    repository = tmp_path / "must-not-backfill"
    git_dir = repository / ".git"
    git_dir.mkdir(parents=True)
    (git_dir / "config").write_text(
        '[remote "origin"]\nurl = https://secret@current.example/private.git\n',
        encoding="utf-8",
    )
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(repository),
        access_mode="read_write",
    )
    later_run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="不要反向补采旧事实",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )

    later_state = await initialize_cowork_state(
        db_session,
        run_id=later_run.id,
        registry=build_default_cowork_registry(),
    )

    assert later_state["session_facts"]["capture_status"] == "legacy_unavailable"
    assert later_state["session_facts"]["workspace_roots"] == []
    assert "current.example" not in json.dumps(later_state["session_facts"])


async def test_cowork_run_api_rejects_ungranted_workspace_file_before_run_creation(
    db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    conversation_id = await ensure_conversation(db_session, title="文件授权边界")
    allowed = tmp_path / "allowed"
    allowed.mkdir()
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(allowed),
        access_mode="read_write",
    )
    outside = tmp_path / "outside.docx"
    outside.write_bytes(b"outside")
    await db_session.commit()
    queue = RecordingCoworkQueue()
    bus = InMemoryRunBus()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    settings = get_settings().model_copy(
        update={"cowork_default_workspace_path": tmp_path / "WorkPilot"}
    )
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://test"
    ) as client:
        response = await client.post(
            "/api/v1/runs/cowork",
            json={
                "conversation_id": str(conversation_id),
                "goal": "修改未授权文件",
                "workspace_files": [str(outside)],
            },
        )

    assert response.status_code == 422
    assert "未获得目录授权" in response.json()["detail"]
    assert queue.run_ids == []
    assert store_sql("SELECT COUNT(*) AS count FROM agent_runs")[0]["count"] == 0


async def test_cowork_run_api_does_not_require_workspace_for_default_permissions(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="默认权限")
    await db_session.commit()
    queue = RecordingCoworkQueue()
    bus = InMemoryRunBus()

    async def override_session() -> AsyncIterator[AsyncSession]:
        yield db_session

    app = create_app()
    default_workspace = tmp_path / "WorkPilot"
    settings = get_settings().model_copy(
        update={"cowork_default_workspace_path": default_workspace}
    )
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[get_run_queue_dependency] = lambda: queue
    app.dependency_overrides[get_run_bus] = lambda: bus
    app.dependency_overrides[require_owner_identity] = lambda: None
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            "/api/v1/runs/cowork",
            json={"conversation_id": str(conversation_id), "goal": "解释什么是 RAG"},
        )

    assert response.status_code == 202
    run_id = UUID(response.json()["run_id"])
    assert queue.run_ids == [run_id]
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run_id)
    assert checkpoint is not None
    assert checkpoint.state["goal"] == "解释什么是 RAG"
    roots = await list_session_roots(db_session, conversation_id=conversation_id)
    assert roots[0].canonical_path == str(default_workspace)
    assert roots[0].label == "WorkPilot 默认文件夹"


async def test_expired_cowork_run_is_recovered_to_cowork_queue_class(
    db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Cowork recovery")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="恢复 Cowork",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=build_default_cowork_registry(),
    )
    store_sql(
        """UPDATE agent_runs
           SET status = 'executing', worker_id = 'lost-worker', lease_until = ?
           WHERE id = ?""",
        (iso_ago(1), str(run.id)),
    )
    await db_session.commit()

    reaped = await reap_expired_runs(db_session)
    await db_session.commit()

    assert reaped.failed == []
    # answer / literature_review 退役后只剩这一个恢复类目。
    assert reaped.recovered_cowork == [(run.id, 1)]


async def test_expired_run_with_open_model_attempt_is_not_replayed(
    db_session: AsyncSession, tmp_path: Path, store_sql
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Unknown model outcome")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="do not replay",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await initialize_cowork_state(
        db_session,
        run_id=run.id,
        registry=build_default_cowork_registry(),
    )
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    await cowork_store().append_session_record(
        run_id=run.id,
        kind="step_attempt",
        operation_id="lost-model-call",
        phase="started",
        payload={
            "source_checkpoint_id": checkpoint.checkpoint_id,
            "result_checkpoint_id": "preallocated-result",
            "iteration": 0,
            "attempt_no": 1,
        },
    )
    store_sql(
        """UPDATE agent_runs
           SET status = 'executing', worker_id = 'lost-worker', lease_until = ?
           WHERE id = ?""",
        (iso_ago(1), str(run.id)),
    )

    reaped = await reap_expired_runs(db_session)
    failed = await get_run(db_session, run.id)

    assert reaped.recovered_cowork == []
    assert reaped.failed == [run.id]
    assert failed is not None and failed.status == "failed"
    assert "outcome unknown" in (failed.error or "")


async def test_completed_model_attempt_is_recovered_without_provider_replay(
    db_session: AsyncSession, tmp_path: Path
) -> None:
    conversation_id = await ensure_conversation(db_session, title="Durable model result")
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="recover paid result",
        budget_tokens=10_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry)
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    assert checkpoint is not None
    identity = {
        "source_checkpoint_id": checkpoint.checkpoint_id,
        "result_checkpoint_id": "durable-model-result-checkpoint",
        "iteration": 0,
        "attempt_no": 1,
    }
    store = cowork_store()
    await store.append_session_record(
        run_id=run.id,
        kind="step_attempt",
        operation_id="durable-model-call",
        phase="started",
        payload=identity,
    )
    await store.append_session_record(
        run_id=run.id,
        kind="step_attempt",
        operation_id="durable-model-call",
        phase="completed",
        payload={
            **identity,
            "result": {
                "stop_reason": "complete",
                "error": None,
                "completion": _completion_record_payload(_final_completion("已从持久记录恢复。")),
            },
        },
    )
    provider = NativeToolProvider([])

    await cowork_run(
        {
            "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
            "session_factory": session_factory,
            "bus": InMemoryRunBus(),
            "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
            "cowork_registry": registry,
        },
        str(run.id),
    )

    finished = await get_run(db_session, run.id)
    assert finished is not None and finished.status == "done"
    assert provider.tool_calls == 0
    saved = await store.load_latest_checkpoint(run_id=run.id)
    assert saved is not None and saved.checkpoint_id == "durable-model-result-checkpoint"


class _SimulatedProcessCrash(BaseException):
    pass


def _crash_model_attempt_at(stage: str):
    def configure(bus: CoworkHookBus) -> None:
        async def crash(context: ModelAttemptHookContext) -> None:
            if context.stage == stage:
                raise _SimulatedProcessCrash(stage)

        bus.model_attempt.register(f"crash-{stage}", crash)

    return configure


async def _initialized_crash_point_run(
    db_session: AsyncSession,
    tmp_path: Path,
    *,
    title: str,
):
    conversation_id = await ensure_conversation(db_session, title=title)
    await create_session_root(
        db_session,
        conversation_id=conversation_id,
        requested_path=str(tmp_path),
        access_mode="read_write",
    )
    run = await create_run(
        db_session,
        conversation_id=conversation_id,
        goal="return one durable answer",
        budget_tokens=100_000,
        budget_calls=10,
        budget_wall_ms=60_000,
        workflow_type="cowork",
    )
    await append_message(
        db_session,
        conversation_id=conversation_id,
        role="user",
        content=run.goal,
        run_id=run.id,
    )
    registry = build_default_cowork_registry()
    await initialize_cowork_state(db_session, run_id=run.id, registry=registry)
    return run, registry


@pytest.mark.parametrize(
    ("stage", "calls_before_recovery", "recoverable", "expected_phases"),
    [
        ("before_started", 0, True, []),
        ("after_started", 0, False, ["started"]),
        ("after_invocation", 1, False, ["started"]),
        ("after_terminal", 1, True, ["started", "completed"]),
    ],
)
async def test_model_attempt_crash_points_never_replay_an_ambiguous_paid_call(
    db_session: AsyncSession,
    tmp_path: Path,
    store_sql,
    stage: str,
    calls_before_recovery: int,
    recoverable: bool,
    expected_phases: list[str],
) -> None:
    run, registry = await _initialized_crash_point_run(
        db_session,
        tmp_path,
        title=f"Crash point {stage}",
    )
    provider = NativeToolProvider([_final_completion("durable answer")])
    context = {
        "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
        "session_factory": session_factory,
        "bus": InMemoryRunBus(),
        "cowork_gateway": ModelGateway(provider, embedding_dimensions=1024),
        "cowork_registry": registry,
        "cowork_hook_configurators": [_crash_model_attempt_at(stage)],
    }

    with pytest.raises(_SimulatedProcessCrash, match=stage):
        await cowork_run(context, str(run.id))

    assert provider.tool_calls == calls_before_recovery
    records = await cowork_store().list_session_records(run_id=run.id)
    assert [record.phase for record in records if record.kind == "step_attempt"] == expected_phases

    store_sql(
        "UPDATE agent_runs SET lease_until = ? WHERE id = ?",
        (iso_ago(1), str(run.id)),
    )
    reaped = await reap_expired_runs(db_session)

    if not recoverable:
        failed = await get_run(db_session, run.id)
        assert reaped.recovered_cowork == []
        assert reaped.failed == [run.id]
        assert failed is not None and failed.status == "failed"
        assert "outcome unknown" in (failed.error or "")
        assert provider.tool_calls == calls_before_recovery
        return

    assert reaped.failed == []
    assert reaped.recovered_cowork == [(run.id, 1)]
    context["cowork_hook_configurators"] = []
    await cowork_run(context, str(run.id))

    finished = await get_run(db_session, run.id)
    assert finished is not None and finished.status == "done"
    # before_started dispatches once during recovery; after_terminal reuses the durable result.
    assert provider.tool_calls == 1


async def test_partial_sse_never_becomes_a_success_checkpoint(
    db_session: AsyncSession,
    tmp_path: Path,
) -> None:
    run, registry = await _initialized_crash_point_run(
        db_session,
        tmp_path,
        title="Partial SSE transport failure",
    )
    requests = 0

    async def handler(request: httpx.Request) -> httpx.Response:
        nonlocal requests
        requests += 1
        return httpx.Response(
            200,
            text=(
                'data: {"model":"served-chat","choices":'
                '[{"index":0,"delta":{"content":"partial answer was cut"},'
                '"finish_reason":"stop"}],"usage":'
                '{"prompt_tokens":5,"completion_tokens":4}}\n\n'
            ),
            headers={"content-type": "text/event-stream"},
        )

    client = httpx.AsyncClient(
        base_url="http://model.test/v1", transport=httpx.MockTransport(handler)
    )
    provider = OpenAICompatibleProvider(
        base_url="http://unused.test/v1",
        api_key="secret",
        chat_model="chat",
        embedding_model="embed",
        client=client,
    )
    try:
        await cowork_run(
            {
                "settings": get_settings().model_copy(update={"run_heartbeat_s": 60.0}),
                "session_factory": session_factory,
                "bus": InMemoryRunBus(),
                "cowork_gateway": ModelGateway(
                    provider,
                    embedding_dimensions=1024,
                    provider_max_retries=0,
                ),
                "cowork_registry": registry,
            },
            str(run.id),
        )
    finally:
        await client.aclose()

    finished = await get_run(db_session, run.id)
    checkpoint = await load_cowork_checkpoint(db_session, run_id=run.id)
    records = await cowork_store().list_session_records(run_id=run.id)
    assert requests == 1
    assert finished is not None and finished.status == "failed"
    assert checkpoint is not None and checkpoint.state["status"] == "failed"
    assert not any(
        message.get("role") == "assistant" and "partial answer" in message.get("content", "")
        for message in checkpoint.state["messages"]
    )
    assert [record.phase for record in records if record.kind == "step_attempt"] == [
        "started",
        "failed",
    ]

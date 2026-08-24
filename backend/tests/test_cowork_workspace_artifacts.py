from pathlib import Path

from docx import Document

from app.cowork.workspace_artifacts import (
    discover_workspace_artifacts,
    snapshot_workspace_artifacts,
)


async def test_discovers_new_and_changed_valid_office_artifacts(tmp_path: Path) -> None:
    existing = tmp_path / "existing.docx"
    document = Document()
    document.add_paragraph("旧内容")
    document.save(existing)
    before = await snapshot_workspace_artifacts(
        tmp_path,
        max_scan_entries=100,
        max_files=20,
    )

    document = Document(existing)
    document.paragraphs[0].text = "新内容"
    document.save(existing)
    created = tmp_path / "created.docx"
    created_document = Document()
    created_document.add_paragraph("新文件")
    created_document.save(created)

    discovery = await discover_workspace_artifacts(
        tmp_path,
        before=before,
        max_scan_entries=100,
        max_files=20,
        max_file_bytes=2 * 1024 * 1024,
    )

    assert {item.path for item in discovery.artifacts} == {existing.resolve(), created.resolve()}
    assert all(item.mime_type.endswith("wordprocessingml.document") for item in discovery.artifacts)
    changed = next(item for item in discovery.artifacts if item.path == existing.resolve())
    assert changed.diff["available"] is True
    assert "-旧内容" in str(changed.diff["text"])
    assert "+新内容" in str(changed.diff["text"])
    added = next(item for item in discovery.artifacts if item.path == created.resolve())
    assert added.diff["created"] is True
    assert "+新文件" in str(added.diff["text"])
    assert discovery.warnings == ()
    assert discovery.truncated is False


async def test_invalid_ooxml_is_not_registered_and_reports_warning(tmp_path: Path) -> None:
    before = await snapshot_workspace_artifacts(
        tmp_path,
        max_scan_entries=100,
        max_files=20,
    )
    (tmp_path / "broken.xlsx").write_bytes(b"not-an-ooxml-file")

    discovery = await discover_workspace_artifacts(
        tmp_path,
        before=before,
        max_scan_entries=100,
        max_files=20,
        max_file_bytes=2 * 1024 * 1024,
    )

    assert discovery.artifacts == ()
    assert len(discovery.warnings) == 1
    assert discovery.warnings[0].startswith("broken.xlsx:")


async def test_broken_before_snapshot_only_disables_diff_not_valid_artifact(
    tmp_path: Path,
) -> None:
    repaired = tmp_path / "repaired.docx"
    repaired.write_bytes(b"not-an-ooxml-file")
    before = await snapshot_workspace_artifacts(
        tmp_path,
        max_scan_entries=100,
        max_files=20,
    )
    document = Document()
    document.add_paragraph("已修复")
    document.save(repaired)

    discovery = await discover_workspace_artifacts(
        tmp_path,
        before=before,
        max_scan_entries=100,
        max_files=20,
        max_file_bytes=2 * 1024 * 1024,
    )

    assert len(discovery.artifacts) == 1
    assert discovery.artifacts[0].path == repaired.resolve()
    assert discovery.artifacts[0].diff["available"] is False
    assert "无法提取差异" in str(discovery.artifacts[0].diff["reason"])
    assert discovery.warnings == ()


async def test_snapshot_skips_hidden_directories_and_symlinks(tmp_path: Path) -> None:
    hidden = tmp_path / ".hidden"
    hidden.mkdir()
    hidden_document = Document()
    hidden_document.add_paragraph("隐藏")
    hidden_document.save(hidden / "hidden.docx")
    hidden_document.save(tmp_path / ".workpilot-backup-hidden.docx")
    outside = tmp_path.parent / f"{tmp_path.name}-outside.docx"
    outside_document = Document()
    outside_document.add_paragraph("外部")
    outside_document.save(outside)
    (tmp_path / "linked.docx").symlink_to(outside)

    snapshot = await snapshot_workspace_artifacts(
        tmp_path,
        max_scan_entries=100,
        max_files=20,
    )

    assert snapshot.files == {}

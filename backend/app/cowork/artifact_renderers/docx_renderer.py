"""DocumentSpec → DOCX。"""

from __future__ import annotations

from pathlib import Path
from unicodedata import east_asian_width

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from app.cowork.artifact_renderers.contracts import DocumentBlock, DocumentSpec
from app.cowork.artifact_renderers.image_assets import compatible_raster_path, image_dimensions


def _set_cell_shading(cell: object, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_style_font(style: object, *, latin: str, east_asia: str, size: float) -> None:
    style.font.name = latin  # type: ignore[attr-defined]
    style.font.size = Pt(size)  # type: ignore[attr-defined]
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()  # type: ignore[attr-defined]
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _set_run_font(run: object, *, latin: str, east_asia: str) -> None:
    run.font.name = latin  # type: ignore[attr-defined]
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()  # type: ignore[attr-defined]
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _display_width(value: object) -> int:
    text = "" if value is None else str(value)
    return sum(2 if east_asian_width(character) in {"W", "F", "A"} else 1 for character in text)


def _set_cell_width(cell: object, width_twips: int) -> None:
    cell.width = Twips(width_twips)  # type: ignore[attr-defined]
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    tc_width = tc_pr.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(width_twips))
    tc_width.set(qn("w:type"), "dxa")


def _set_cell_margins(
    cell: object, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row: object) -> None:
    properties = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _configure(document: DocumentObject) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, latin="Arial", east_asia="Microsoft YaHei", size=10.5)
    normal.font.color.rgb = RGBColor(35, 47, 42)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True
    for level, size in ((1, 20), (2, 15), (3, 12)):
        style = styles[f"Heading {level}"]
        _set_style_font(style, latin="Arial", east_asia="Microsoft YaHei", size=size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(22, 122, 91)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
    for name, size in (("Title", 28), ("Subtitle", 13), ("Caption", 9)):
        _set_style_font(
            styles[name],
            latin="Arial",
            east_asia="Microsoft YaHei",
            size=size,
        )
    styles["Title"].paragraph_format.keep_with_next = True
    styles["Subtitle"].paragraph_format.keep_with_next = True
    list_style = styles["List Bullet"]
    _set_style_font(list_style, latin="Arial", east_asia="Microsoft YaHei", size=10.5)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.15)
    list_style.paragraph_format.space_after = Pt(4)
    if "WorkPilot Callout" not in styles:
        callout = styles.add_style("WorkPilot Callout", WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(callout, latin="Arial", east_asia="Microsoft YaHei", size=10.5)
        callout.font.bold = True
        callout.font.color.rgb = RGBColor(42, 65, 56)


def _add_table(document: DocumentObject, block: DocumentBlock) -> None:
    column_count = len(block.headers) or max((len(row) for row in block.rows), default=1)
    table = document.add_table(rows=1 if block.headers else 0, cols=column_count)
    table.style = "Light Shading Accent 1"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_properties = table._tbl.tblPr
    layout = table_properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    values_by_column: list[list[object]] = [
        [block.headers[index] if index < len(block.headers) else ""]
        + [row[index] if index < len(row) else "" for row in block.rows[:100]]
        for index in range(column_count)
    ]
    weights = [
        max(6, min(36, max((_display_width(value) for value in values), default=6)))
        for values in values_by_column
    ]
    section = document.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    if page_width is None or left_margin is None or right_margin is None:  # pragma: no cover
        raise ValueError("DOCX 页面尺寸或页边距缺失")
    available_twips = int(page_width.twips - left_margin.twips - right_margin.twips)
    raw_widths = [max(720, round(available_twips * weight / sum(weights))) for weight in weights]
    scale = available_twips / sum(raw_widths)
    widths = [max(600, round(width * scale)) for width in raw_widths]
    widths[-1] += available_twips - sum(widths)
    if block.headers:
        for index, value in enumerate(block.headers):
            cell = table.rows[0].cells[index]
            cell.text = value
            _set_cell_shading(cell, "DCEBE5")
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        _repeat_table_header(table.rows[0])
    for values in block.rows:
        cells = table.add_row().cells
        for index in range(column_count):
            cell_value = values[index] if index < len(values) else None
            cells[index].text = "" if cell_value is None else str(cell_value)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_together = True


def _add_image(document: DocumentObject, block: DocumentBlock) -> None:
    if block.image_path is None:  # guarded by DocumentBlock validation
        raise ValueError("image block 缺少 image_path")
    source = Path(block.image_path)
    source_width, source_height = image_dimensions(source)
    ratio = source_width / source_height
    width_inches = block.image_width_inches or 6.2
    width_inches = min(width_inches, 7.2 * ratio)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = bool(block.image_caption)
    with compatible_raster_path(source, max_dimension=1800) as compatible_path:
        inline = paragraph.add_run().add_picture(
            str(compatible_path),
            width=Inches(width_inches),
        )
    inline._inline.docPr.set("descr", block.image_alt or "")
    if block.image_caption:
        caption = document.add_paragraph(block.image_caption, style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_with_next = False


def _add_block(document: DocumentObject, block: DocumentBlock) -> None:
    if block.type == "paragraph":
        paragraph = document.add_paragraph(block.text or "")
        if block.style == "lead":
            paragraph.style = document.styles["Subtitle"]
        elif block.style == "caption":
            paragraph.style = document.styles["Caption"]
    elif block.type == "bullets":
        for value in block.items:
            document.add_paragraph(value, style="List Bullet")
    elif block.type == "table":
        _add_table(document, block)
    elif block.type == "image":
        _add_image(document, block)
    elif block.type == "quote":
        paragraph = document.add_paragraph(block.text or "", style="Intense Quote")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif block.type == "callout":
        paragraph = document.add_paragraph(block.text or "", style="WorkPilot Callout")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(10)


def render_document(spec: DocumentSpec, target: Path) -> None:
    document: DocumentObject = Document()
    _configure(document)
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(spec.title)
    _set_run_font(run, latin="Arial", east_asia="Microsoft YaHei")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 33, 29)
    if spec.subtitle:
        document.add_paragraph(spec.subtitle, style="Subtitle")
    if spec.author:
        author = document.add_paragraph(spec.author)
        author.style = document.styles["Caption"]
    for section in spec.sections:
        document.add_heading(section.heading, level=section.level)
        for block in section.blocks:
            _add_block(document, block)
    document.core_properties.title = spec.title
    document.core_properties.subject = spec.purpose or ""
    document.core_properties.author = spec.author or "WorkPilot"
    document.save(str(target))


__all__ = ["render_document"]
